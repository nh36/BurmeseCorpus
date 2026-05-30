from __future__ import annotations

import json
import os
import subprocess
import sys
import tempfile
import unittest
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
SCRIPTS_DIR = ROOT / "scripts"
if str(SCRIPTS_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPTS_DIR))

from bibtex_common import parse_bibtex_text
from build_bibtex_authority import (
    ACRONYM_STATUS_FIELDS,
    AUTHORITY_FIELDS,
    CROSSWALK_FIELDS,
    REMAINING_ACRONYMS,
    SOURCE_FAMILY_FIELDS,
    build_ippa_review_artifacts,
    build_authority,
    parse_locator,
)
from corpus_common import read_tsv, write_tsv
from extract_bibliography_acronyms import PRIORITY_ACRONYMS, extract_documentation_sections, extract_explicit_definition_candidates
from extract_frasch_bibliography import run_extraction
from harvest_local_bibliography_sources import run_harvest
from import_external_bibtex import import_external_bibtex
from validate_bibtex_authority import validate_bibtex_authority


FAMILY_FIELDS = [
    "family_id",
    "family_label",
    "family_type",
    "member_count",
    "occurrence_count",
    "sample_raw_references",
    "likely_contains_translation",
    "review_status",
    "notes",
]

MEMBER_FIELDS = [
    "family_id",
    "raw_reference_string",
    "occurrence_count",
    "example_record_ids",
    "notes",
]

CANDIDATE_FIELDS = [
    "work_candidate_id",
    "family_id",
    "provisional_short_label",
    "author_original",
    "author_normalized",
    "year",
    "title_original",
    "title_normalized",
    "publication_details",
    "language",
    "script",
    "translation_relevance",
    "evidence_raw_references",
    "review_status",
    "notes",
]

SEED_FIELDS = [
    "abbreviation",
    "family_id",
    "provisional_label",
    "probable_bibtex_key",
    "source_type",
    "confidence",
    "evidence",
    "needs_human_review",
    "notes",
]

LOCAL_CANDIDATE_FIELDS = [
    "candidate_id",
    "search_term",
    "name",
    "original_path",
    "file_type",
    "file_size",
    "sha256",
    "probable_work_label",
    "probable_author",
    "probable_year",
    "match_confidence",
    "copy_status",
    "copied_path",
    "notes",
]

LOCAL_MANIFEST_FIELDS = [
    "source_file_id",
    "original_path",
    "copied_path",
    "file_name",
    "file_type",
    "file_size",
    "sha256",
    "source_folder_hint",
    "copy_date",
    "copy_status",
    "notes",
]


class BibtexAuthorityTests(unittest.TestCase):
    def test_import_external_bibtex_preserves_entries_and_warnings(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            input_path = temp_path / "sample.bib"
            input_path.write_text(
                "@book{harvey1925,\n"
                "  author = {Harvey, G. E.},\n"
                "  title = {History of Burma},\n"
                "  year = {1925}\n"
                "}\n\n"
                "{brokenKey,\n"
                "  title = {Broken but preserved}\n"
                "}\n",
                encoding="utf-8",
            )
            output_dir = temp_path / "external"
            report = import_external_bibtex(input_path, "sample.bib", output_dir)
            entries_path = output_dir / "sample_entries.tsv"
            report_path = output_dir / "sample_import_report.json"

            self.assertEqual(report["entry_count"], 2)
            self.assertTrue(report["parse_warnings"])
            self.assertTrue(entries_path.exists())
            self.assertTrue(report_path.exists())

            rows = entries_path.read_text(encoding="utf-8")
            self.assertIn("harvey1925", rows)
            self.assertIn("brokenKey", rows)
            self.assertIn("raw_entry_hash", rows)

    def test_parser_salvages_raw_bibtex_without_type(self) -> None:
        entries, warnings = parse_bibtex_text("{brokenKey,\n  title = {Broken}\n}\n", source_label="test")
        self.assertEqual(len(entries), 1)
        self.assertEqual(entries[0]["entry_type"], "unknown")
        self.assertTrue(any("salvaged malformed entry" in warning for warning in warnings))

    def test_locator_parsing_examples(self) -> None:
        self.assertEqual(parse_locator("OBI 3, p. 2", "fam-obi-internal", "OBI"), ("3, p. 2", "volume_page"))
        self.assertEqual(parse_locator("Pl. II 198", "fam-plate-references", "Pl."), ("II 198", "plate"))
        self.assertEqual(parse_locator("List 90", "fam-list-catalogue", "List"), ("90", "catalogue_number"))
        self.assertEqual(parse_locator("IPPA 101-102", "fam-ippa-catalogue", "IPPA references"), ("101-102", "number"))
        self.assertEqual(parse_locator("OR folio 12 verso", "fam-or-catalogue", "OR references"), ("folio 12 verso", "folio"))
        self.assertEqual(parse_locator("RDASB 1971", "fam-rdasb-publication", "RDASB"), ("1971", "year"))

    def test_extract_explicit_definition_candidates_reads_abbreviation_list(self) -> None:
        rows, rejected = extract_explicit_definition_candidates(
            "PPA = Inscriptions of Pagan, Pinya and Ava\nRDASB = Report of the Director, Archaeological Survey of Burma\n",
            source_file_id="doc-1",
            source_file_label="mock list",
            acronyms=["PPA", "RDASB"],
        )
        by_acronym = {row["acronym"]: row for row in rows}
        self.assertEqual(by_acronym["PPA"]["candidate_expansion"], "Inscriptions of Pagan, Pinya and Ava")
        self.assertEqual(by_acronym["PPA"]["evidence_type"], "explicit_abbreviation_list")
        self.assertEqual(by_acronym["RDASB"]["definition_quality"], "explicit")
        self.assertEqual(rejected, [])

    def test_extract_explicit_definition_candidates_ignores_contextual_usage(self) -> None:
        rows, rejected = extract_explicit_definition_candidates(
            "PPA, p. 55\nPl. II 198\n",
            source_file_id="doc-2",
            source_file_label="mock contexts",
            acronyms=["PPA", "Pl."],
        )
        self.assertEqual(rows, [])
        self.assertEqual(rejected, [])

    def test_extract_explicit_definition_candidates_rejects_known_false_positives(self) -> None:
        rows, rejected = extract_explicit_definition_candidates(
            "Date: CS 581 (List)\nspelling of inscription (OBI)\nor: Palm-leaf manuscript\n",
            source_file_id="doc-3",
            source_file_label="mock bad cases",
            acronyms=["List", "OBI", "OR"],
        )
        self.assertEqual(rows, [])
        reasons = {row["acronym"]: row["reason_rejected"] for row in rejected}
        self.assertIn("candidate expansion does not look like a title or source name", reasons["List"])
        self.assertIn("candidate expansion does not look like a title or source name", reasons["OBI"])
        self.assertIn("lowercase 'or' is not the OR acronym", reasons["OR"])

    def test_extract_explicit_definition_candidates_keeps_multiline_sip_definition(self) -> None:
        rows, rejected = extract_explicit_definition_candidates(
            "SIP = Pe Maung Tin and G. H. Luce, Selections from the\nInscriptions of Pagan\n",
            source_file_id="doc-4",
            source_file_label="mock sip list",
            acronyms=["SIP"],
        )
        self.assertEqual(rejected, [])
        self.assertEqual(len(rows), 1)
        self.assertEqual(
            rows[0]["candidate_expansion"],
            "Pe Maung Tin and G. H. Luce, Selections from the Inscriptions of Pagan",
        )
        self.assertEqual(rows[0]["definition_quality"], "explicit")

    def write_fixture_tables(self, base: Path) -> tuple[Path, Path, Path, Path, Path]:
        families_path = base / "reference_families.tsv"
        members_path = base / "reference_family_members.tsv"
        candidates_path = base / "bibliographic_work_candidates.tsv"
        seeds_path = base / "source_abbreviation_seeds.tsv"
        external_entries_path = base / "asia_2_entries.tsv"

        write_tsv(
            families_path,
            [
                {
                    "family_id": "fam-obi-internal",
                    "family_label": "OBI",
                    "family_type": "source_catalogue",
                    "member_count": "1",
                    "occurrence_count": "1",
                    "sample_raw_references": "OBI 3, p. 2",
                    "likely_contains_translation": "no",
                    "review_status": "reviewed_provisional",
                    "notes": "",
                },
                {
                    "family_id": "fam-plate-references",
                    "family_label": "Pl.",
                    "family_type": "internal_reference",
                    "member_count": "1",
                    "occurrence_count": "1",
                    "sample_raw_references": "Pl. II 198",
                    "likely_contains_translation": "no",
                    "review_status": "needs_human_review",
                    "notes": "",
                },
                {
                    "family_id": "fam-list-catalogue",
                    "family_label": "List",
                    "family_type": "source_catalogue",
                    "member_count": "1",
                    "occurrence_count": "1",
                    "sample_raw_references": "List 90",
                    "likely_contains_translation": "no",
                    "review_status": "needs_human_review",
                    "notes": "",
                },
                {
                    "family_id": "fam-rdasb-publication",
                    "family_label": "RDASB",
                    "family_type": "publication",
                    "member_count": "1",
                    "occurrence_count": "1",
                    "sample_raw_references": "RDASB 1971",
                    "likely_contains_translation": "possible",
                    "review_status": "reviewed_provisional",
                    "notes": "",
                },
                {
                    "family_id": "fam-harvey-history",
                    "family_label": "Harvey, History",
                    "family_type": "book",
                    "member_count": "1",
                    "occurrence_count": "1",
                    "sample_raw_references": "Harvey, History, p. 10",
                    "likely_contains_translation": "possible",
                    "review_status": "reviewed_provisional",
                    "notes": "",
                },
                {
                    "family_id": "fam-unresolved",
                    "family_label": "Mystery Source",
                    "family_type": "book",
                    "member_count": "1",
                    "occurrence_count": "1",
                    "sample_raw_references": "Mystery Source 12",
                    "likely_contains_translation": "unknown",
                    "review_status": "unreviewed",
                    "notes": "",
                },
            ],
            FAMILY_FIELDS,
        )
        write_tsv(
            members_path,
            [
                {"family_id": "fam-obi-internal", "raw_reference_string": "OBI 3, p. 2", "occurrence_count": "2", "example_record_ids": "obi-1", "notes": ""},
                {"family_id": "fam-plate-references", "raw_reference_string": "Pl. II 198", "occurrence_count": "1", "example_record_ids": "obi-2", "notes": ""},
                {"family_id": "fam-list-catalogue", "raw_reference_string": "List 90", "occurrence_count": "3", "example_record_ids": "obi-3", "notes": ""},
                {"family_id": "fam-rdasb-publication", "raw_reference_string": "RDASB 1971", "occurrence_count": "2", "example_record_ids": "obi-4", "notes": ""},
                {"family_id": "fam-harvey-history", "raw_reference_string": "Harvey, History, p. 10", "occurrence_count": "1", "example_record_ids": "obi-5", "notes": ""},
                {"family_id": "fam-unresolved", "raw_reference_string": "Mystery Source 12", "occurrence_count": "1", "example_record_ids": "obi-6", "notes": ""},
            ],
            MEMBER_FIELDS,
        )
        write_tsv(
            candidates_path,
            [
                {
                    "work_candidate_id": "wc-obi",
                    "family_id": "fam-obi-internal",
                    "provisional_short_label": "OBI",
                    "author_original": "",
                    "author_normalized": "",
                    "year": "",
                    "title_original": "Old Burmese Inscriptions Corpus",
                    "title_normalized": "old burmese inscriptions corpus",
                    "publication_details": "",
                    "language": "my",
                    "script": "Mymr",
                    "translation_relevance": "unlikely_translation",
                    "evidence_raw_references": "OBI 3, p. 2",
                    "review_status": "reviewed_provisional",
                    "notes": "",
                },
                {
                    "work_candidate_id": "wc-list",
                    "family_id": "fam-list-catalogue",
                    "provisional_short_label": "List",
                    "author_original": "",
                    "author_normalized": "",
                    "year": "",
                    "title_original": "List catalogue reference",
                    "title_normalized": "list catalogue reference",
                    "publication_details": "",
                    "language": "",
                    "script": "",
                    "translation_relevance": "unknown",
                    "evidence_raw_references": "List 90",
                    "review_status": "needs_human_review",
                    "notes": "",
                },
                {
                    "work_candidate_id": "wc-rdasb",
                    "family_id": "fam-rdasb-publication",
                    "provisional_short_label": "RDASB",
                    "author_original": "",
                    "author_normalized": "",
                    "year": "",
                    "title_original": "Report of the Director, Archaeological Survey of Burma",
                    "title_normalized": "report of the director archaeological survey of burma",
                    "publication_details": "",
                    "language": "",
                    "script": "",
                    "translation_relevance": "possible_translation",
                    "evidence_raw_references": "RDASB 1971",
                    "review_status": "reviewed_provisional",
                    "notes": "",
                },
                {
                    "work_candidate_id": "wc-harvey",
                    "family_id": "fam-harvey-history",
                    "provisional_short_label": "Harvey, History",
                    "author_original": "Harvey",
                    "author_normalized": "harvey",
                    "year": "1925",
                    "title_original": "History of Burma",
                    "title_normalized": "history of burma",
                    "publication_details": "",
                    "language": "",
                    "script": "",
                    "translation_relevance": "possible_translation",
                    "evidence_raw_references": "Harvey, History, p. 10",
                    "review_status": "reviewed_provisional",
                    "notes": "",
                },
                {
                    "work_candidate_id": "wc-unresolved",
                    "family_id": "fam-unresolved",
                    "provisional_short_label": "Mystery Source",
                    "author_original": "Anon.",
                    "author_normalized": "",
                    "year": "",
                    "title_original": "Mystery Chronicle of Bagan",
                    "title_normalized": "mystery chronicle of bagan",
                    "publication_details": "",
                    "language": "",
                    "script": "",
                    "translation_relevance": "unknown",
                    "evidence_raw_references": "Mystery Source 12",
                    "review_status": "unreviewed",
                    "notes": "",
                },
            ],
            CANDIDATE_FIELDS,
        )
        write_tsv(
            seeds_path,
            [
                {
                    "abbreviation": "OBI",
                    "family_id": "fam-obi-internal",
                    "provisional_label": "Old Burmese Inscriptions corpus source set",
                    "probable_bibtex_key": "obiCorpusSource",
                    "source_type": "source_catalogue",
                    "confidence": "high",
                    "evidence": "fixture seed",
                    "needs_human_review": "false",
                    "notes": "",
                },
                {
                    "abbreviation": "List",
                    "family_id": "fam-list-catalogue",
                    "provisional_label": "List catalogue",
                    "probable_bibtex_key": "listCatalogueUnresolved",
                    "source_type": "source_catalogue",
                    "confidence": "medium",
                    "evidence": "triage seed",
                    "needs_human_review": "true",
                    "notes": "",
                }
            ],
            SEED_FIELDS,
        )
        write_tsv(
            external_entries_path,
            [
                {
                    "bibtex_key": "harveyExt",
                    "entry_type": "book",
                    "author": "Harvey, G. E.",
                    "editor": "",
                    "year": "1925",
                    "title": "History of Burma",
                    "booktitle": "",
                    "journal": "",
                    "publisher": "Longmans",
                    "address": "London",
                    "doi": "",
                    "url": "",
                    "isbn": "",
                    "raw_entry_hash": "abc",
                    "source_label": "asia 2.bib",
                    "notes": "",
                }
            ],
            [
                "bibtex_key",
                "entry_type",
                "author",
                "editor",
                "year",
                "title",
                "booktitle",
                "journal",
                "publisher",
                "address",
                "doi",
                "url",
                "isbn",
                "raw_entry_hash",
                "source_label",
                "notes",
            ],
        )
        return families_path, members_path, candidates_path, seeds_path, external_entries_path

    def write_fixture_acronym_files(self, base: Path) -> tuple[Path, Path]:
        candidates_path = base / "acronym_definition_candidates.tsv"
        report_path = base / "acronym_definition_report.json"
        write_tsv(
            candidates_path,
            [
                {
                    "candidate_id": "doc:ppa",
                    "acronym": "PPA",
                    "candidate_expansion": "Inscriptions of Pagan, Pinya and Ava",
                    "raw_definition": "Inscriptions of Pagan, Pinya and Ava (PPA)",
                    "definition_context": "Inscriptions of Pagan, Pinya and Ava (PPA) was the first volume.",
                    "source_file_id": "doc-1",
                    "source_file_label": "Frasch translation",
                    "source_location_hint": "pattern hit",
                    "evidence_type": "explicit_parenthetical_definition",
                    "confidence": "high",
                    "definition_quality": "explicit",
                    "needs_human_review": "false",
                    "notes": "",
                },
                {
                    "candidate_id": "doc:ub",
                    "acronym": "UB",
                    "candidate_expansion": "Inscriptions Collected in Upper Burma",
                    "raw_definition": "Inscriptions collected in Upper Burma (UB 1, UB 2)",
                    "definition_context": "Inscriptions collected in Upper Burma (UB 1, UB 2).",
                    "source_file_id": "doc-1",
                    "source_file_label": "Frasch translation",
                    "source_location_hint": "pattern hit",
                    "evidence_type": "explicit_parenthetical_definition",
                    "confidence": "high",
                    "definition_quality": "explicit",
                    "needs_human_review": "false",
                    "notes": "",
                },
                {
                    "candidate_id": "doc:a",
                    "acronym": "A",
                    "candidate_expansion": "Original Inscriptions Collected by King Bodawpaya and now placed near the Patodawgyi Pagoda, Amarapura",
                    "raw_definition": "Original Inscriptions Collected by King Bodawpaya and now placed near the Patodawgyi Pagoda, Amarapura (A)",
                    "definition_context": "Original Inscriptions Collected by King Bodawpaya and now placed near the Patodawgyi Pagoda, Amarapura (A).",
                    "source_file_id": "doc-1",
                    "source_file_label": "Frasch translation",
                    "source_location_hint": "pattern hit",
                    "evidence_type": "explicit_parenthetical_definition",
                    "confidence": "high",
                    "definition_quality": "explicit",
                    "needs_human_review": "false",
                    "notes": "",
                },
                {
                    "candidate_id": "doc:arasi",
                    "acronym": "ARASI",
                    "candidate_expansion": "Annual Reports of the Archaeological Survey of India",
                    "raw_definition": "ARASI = Annual Reports of the Archaeological Survey of India",
                    "definition_context": "ARASI = Annual Reports of the Archaeological Survey of India.",
                    "source_file_id": "doc-1",
                    "source_file_label": "Frasch translation",
                    "source_location_hint": "pattern hit",
                    "evidence_type": "explicit_abbreviation_list",
                    "confidence": "high",
                    "definition_quality": "explicit",
                    "needs_human_review": "false",
                    "notes": "",
                },
                {
                    "candidate_id": "doc:b",
                    "acronym": "B",
                    "candidate_expansion": "Inscriptions Copied from the Stones Collected by King Bodawpaya and Placed near the Aracan Pagoda",
                    "raw_definition": "Inscriptions Copied from the Stones Collected by King Bodawpaya and Placed near the Aracan Pagoda (B 1, B 2)",
                    "definition_context": "Inscriptions Copied from the Stones Collected by King Bodawpaya and Placed near the Aracan Pagoda (B 1, B 2).",
                    "source_file_id": "doc-1",
                    "source_file_label": "Frasch translation",
                    "source_location_hint": "pattern hit",
                    "evidence_type": "explicit_parenthetical_definition",
                    "confidence": "high",
                    "definition_quality": "explicit",
                    "needs_human_review": "false",
                    "notes": "",
                },
                {
                    "candidate_id": "doc:bbhc",
                    "acronym": "BBHC",
                    "candidate_expansion": "Bulletin of the Burma Historical Commission",
                    "raw_definition": "Bulletin of the Burma Historical Commission (BBHC)",
                    "definition_context": "shortened summary in the Bulletin of the Burma Historical Commission (BBHC).",
                    "source_file_id": "doc-1",
                    "source_file_label": "Frasch translation",
                    "source_location_hint": "pattern hit",
                    "evidence_type": "explicit_parenthetical_definition",
                    "confidence": "high",
                    "definition_quality": "explicit",
                    "needs_human_review": "false",
                    "notes": "",
                },
                {
                    "candidate_id": "doc:bedb",
                    "acronym": "BED B",
                    "candidate_expansion": "Bagan Epigraphic Database, Part B",
                    "raw_definition": "Bagan Epigraphic Database (BED) — PART B",
                    "definition_context": "Bagan Epigraphic Database (BED) — PART B",
                    "source_file_id": "doc-2",
                    "source_file_label": "Bagan database",
                    "source_location_hint": "heading",
                    "evidence_type": "explicit_abbreviation_list",
                    "confidence": "high",
                    "definition_quality": "strong",
                    "needs_human_review": "false",
                    "notes": "",
                },
                {
                    "candidate_id": "doc:mm",
                    "acronym": "MM",
                    "candidate_expansion": "Middle Mon",
                    "raw_definition": "MM = Middle Mon",
                    "definition_context": "MM = Middle Mon.",
                    "source_file_id": "doc-1",
                    "source_file_label": "Frasch translation",
                    "source_location_hint": "pattern hit",
                    "evidence_type": "explicit_abbreviation_list",
                    "confidence": "high",
                    "definition_quality": "explicit",
                    "needs_human_review": "false",
                    "notes": "",
                },
                {
                    "candidate_id": "doc:sip",
                    "acronym": "SIP",
                    "candidate_expansion": "Pe Maung Tin and G. H. Luce, Selections from the Inscriptions of Pagan",
                    "raw_definition": "SIP = Pe Maung Tin and G. H. Luce, Selections from the Inscriptions of Pagan",
                    "definition_context": "SIP = Pe Maung Tin and G. H. Luce, Selections from the Inscriptions of Pagan",
                    "source_file_id": "doc-3",
                    "source_file_label": "Luce comparative wordlist",
                    "source_location_hint": "line 12",
                    "evidence_type": "explicit_abbreviation_list",
                    "confidence": "high",
                    "definition_quality": "explicit",
                    "needs_human_review": "false",
                    "notes": "",
                },
                {
                    "candidate_id": "doc:pl",
                    "acronym": "Pl.",
                    "candidate_expansion": "",
                    "raw_definition": "Pl. II 198",
                    "definition_context": "Pl. II 198 remains a plate locator.",
                    "source_file_id": "doc-2",
                    "source_file_label": "Bagan database",
                    "source_location_hint": "context",
                    "evidence_type": "contextual_usage",
                    "confidence": "low",
                    "definition_quality": "context_only",
                    "needs_human_review": "true",
                    "notes": "",
                },
                {
                    "candidate_id": "doc:rdasb",
                    "acronym": "RDASB",
                    "candidate_expansion": "",
                    "raw_definition": "",
                    "definition_context": "",
                    "source_file_id": "",
                    "source_file_label": "",
                    "source_location_hint": "searched corpus docs",
                    "evidence_type": "negative_evidence",
                    "confidence": "low",
                    "definition_quality": "not_found",
                    "needs_human_review": "true",
                    "notes": "",
                },
            ],
            [
                "candidate_id",
                "acronym",
                "candidate_expansion",
                "raw_definition",
                "definition_context",
                "source_file_id",
                "source_file_label",
                "source_location_hint",
                "evidence_type",
                "confidence",
                "definition_quality",
                "needs_human_review",
                "notes",
            ],
        )
        report_path.write_text(
            json.dumps(
                {
                    "documentation_files_searched_count": 2,
                    "frasch_stadt_staat_files_searched_count": 1,
                    "fratsch_stadt_staat_files_searched_count": 1,
                    "bagan_database_context_matches": 1,
                    "ocr_needed_count": 0,
                }
            )
            + "\n",
            encoding="utf-8",
        )
        return candidates_path, report_path

    def write_manual_seed_file(self, base: Path) -> Path:
        manual_seed_path = base / "manual_acronym_seeds.tsv"
        write_tsv(
            manual_seed_path,
            [
                {
                    "acronym": "EB",
                    "expansion": "Epigraphia Birmanica",
                    "authority_key": "epigraphiaBirmanica",
                    "source_family_id": "sf-eb",
                    "confidence": "high",
                    "supplied_by": "Nathan Hill",
                    "date_added": "2026-05-30",
                    "needs_documentary_confirmation": "true",
                    "notes": "Manual seed; seek confirmation in local abbreviation lists.",
                },
                {
                    "acronym": "JBRS",
                    "expansion": "Journal of the Burma Research Society",
                    "authority_key": "journalBurmaResearchSociety",
                    "source_family_id": "sf-jbrs",
                    "confidence": "high",
                    "supplied_by": "Nathan Hill",
                    "date_added": "2026-05-30",
                    "needs_documentary_confirmation": "true",
                    "notes": "Manual seed; seek confirmation in local abbreviation lists.",
                },
                {
                    "acronym": "JRAS",
                    "expansion": "Journal of the Royal Asiatic Society",
                    "authority_key": "journalRoyalAsiaticSociety",
                    "source_family_id": "sf-jras",
                    "confidence": "high",
                    "supplied_by": "Nathan Hill",
                    "date_added": "2026-05-30",
                    "needs_documentary_confirmation": "true",
                    "notes": "Manual seed; seek confirmation in local abbreviation lists.",
                },
                {
                    "acronym": "OBI",
                    "expansion": "Old Burmese Inscriptions",
                    "authority_key": "obiCorpusSource",
                    "source_family_id": "sf-obi",
                    "confidence": "high",
                    "supplied_by": "Nathan Hill",
                    "date_added": "2026-05-30",
                    "needs_documentary_confirmation": "true",
                    "notes": "Manual seed; seek confirmation in corpus documentation.",
                },
            ],
            [
                "acronym",
                "expansion",
                "authority_key",
                "source_family_id",
                "confidence",
                "supplied_by",
                "date_added",
                "needs_documentary_confirmation",
                "notes",
            ],
        )
        return manual_seed_path

    def test_build_authority_generates_seed_authority_and_stub(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            families_path, members_path, candidates_path, seeds_path, external_entries_path = self.write_fixture_tables(temp_path)
            acronym_candidates_path, acronym_report_path = self.write_fixture_acronym_files(temp_path)
            manual_seed_path = self.write_manual_seed_file(temp_path)
            output_dir = temp_path / "authority"

            report = build_authority(
                reference_families_path=families_path,
                reference_members_path=members_path,
                work_candidates_path=candidates_path,
                seed_path=seeds_path,
                external_entries_path=external_entries_path,
                output_dir=output_dir,
                frasch_references_path=temp_path / "missing_frasch.tsv",
                local_candidates_path=temp_path / "missing_local_candidates.tsv",
                local_manifest_path=temp_path / "missing_local_manifest.tsv",
                acronym_candidates_path=acronym_candidates_path,
                acronym_report_path=acronym_report_path,
            )

            authority_bib = (output_dir / "bibliography_authority.bib").read_text(encoding="utf-8")
            candidates_bib = (output_dir / "bibliography_candidates.bib").read_text(encoding="utf-8")
            authority_tsv = (output_dir / "bibtex_authority.tsv").read_text(encoding="utf-8")
            crosswalk_tsv = (output_dir / "raw_reference_to_bibtex.tsv").read_text(encoding="utf-8")

            self.assertIn("obiCorpusSource", authority_bib)
            self.assertIn("duroiselle1921list", authority_bib)
            self.assertIn("Provisional entry generated from corpus reference triage; requires human review", candidates_bib)
            self.assertIn("obiCorpusSource", crosswalk_tsv)
            self.assertIn("source_family_match", crosswalk_tsv)
            self.assertIn("familyid = {fam-unresolved}", candidates_bib)
            self.assertIn("Mystery Source", candidates_bib)
            self.assertGreater(report["authority_entry_count"], 0)
            self.assertGreater(report["candidate_entry_count"], 0)
            self.assertIn("matched_external_key", authority_tsv)
            self.assertTrue((output_dir / "source_family_authority.tsv").exists())
            self.assertTrue((output_dir / "source_work_authority.tsv").exists())
            self.assertTrue((output_dir / "source_work_authority_audit.tsv").exists())
            self.assertTrue((output_dir / "source_work_to_bibtex_reconciliation.tsv").exists())
            self.assertTrue((output_dir / "bibtex_field_quality_audit.tsv").exists())
            self.assertTrue((output_dir / "authority_key_normalization.tsv").exists())

    def test_harvest_local_bibliography_sources_finds_frasch_and_frosch_and_copies(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            author_root = temp_path / "Authors alphabetical"
            frasch_dir = author_root / "Frasch, Tilmans"
            frosch_dir = author_root / "Frosch"
            frasch_dir.mkdir(parents=True)
            frosch_dir.mkdir(parents=True)
            (frasch_dir / "Bagan Epig Database.doc").write_text("frasch dummy", encoding="utf-8")
            (frosch_dir / "Tilman bibliography.rtf").write_text("frosch dummy", encoding="utf-8")
            downloads = temp_path / "Downloads"
            downloads.mkdir()

            previous = {key: os.environ.get(key) for key in ("OBI_AUTHOR_ALPHA_ROOT", "OBI_LIBRARY_ROOT", "OBI_LOCAL_BIB_ROOT")}
            os.environ["OBI_AUTHOR_ALPHA_ROOT"] = str(author_root)
            os.environ["OBI_LIBRARY_ROOT"] = str(temp_path)
            os.environ["OBI_LOCAL_BIB_ROOT"] = str(downloads)
            try:
                output_dir = temp_path / "output"
                report = run_harvest("frasch", output_dir)
            finally:
                for key, value in previous.items():
                    if value is None:
                        os.environ.pop(key, None)
                    else:
                        os.environ[key] = value

            self.assertGreaterEqual(report["raw_candidate_count"], 2)
            candidates = (output_dir / "frasch_source_candidates.tsv").read_text(encoding="utf-8")
            self.assertIn("Bagan Epig Database.doc", candidates)
            self.assertIn("Tilman bibliography.rtf", candidates)
            manifest = (output_dir / "local_file_manifest.tsv").read_text(encoding="utf-8")
            self.assertIn("data/local/bibliography_sources", manifest)

    def test_extract_frasch_bibliography_from_docx(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            document_path = temp_path / "frasch.docx"
            from docx import Document

            document = Document()
            document.add_paragraph("Number: 1")
            document.add_paragraph("References: U Tha Myat, The Pali Version of the Myazedi Inscription, Rangoon 1958; List 90")
            document.save(document_path)

            manifest_path = temp_path / "manifest.tsv"
            write_tsv(
                manifest_path,
                [
                    {
                        "source_file_id": "frasch-docx",
                        "original_path": "OBI_AUTHOR_ALPHA_ROOT:Frasch/frasch.docx",
                        "copied_path": document_path.as_posix(),
                        "file_name": "frasch.docx",
                        "file_type": "docx",
                        "file_size": "1",
                        "sha256": "abc",
                        "source_folder_hint": "Frasch",
                        "copy_date": "2024-01-01T00:00:00Z",
                        "copy_status": "copied",
                        "notes": "",
                    }
                ],
                LOCAL_MANIFEST_FIELDS,
            )
            output_dir = temp_path / "output"
            report = run_extraction(manifest_path, output_dir)
            rows = (output_dir / "frasch_reference_entries.tsv").read_text(encoding="utf-8")
            self.assertEqual(report["source_file_count"], 1)
            self.assertIn("The Pali Version of the Myazedi Inscription", rows)
            self.assertIn("List 90", rows)
            self.assertIn("\tlow\t", rows)
            self.assertIn("\thigh\t", rows)
            bibliography = (output_dir / "frasch_bibliography.bib").read_text(encoding="utf-8")
            self.assertIn("U Tha Myat", bibliography)

    def test_build_authority_promotes_manual_local_match(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            families_path = temp_path / "families.tsv"
            members_path = temp_path / "members.tsv"
            candidates_path = temp_path / "candidates.tsv"
            seeds_path = temp_path / "seeds.tsv"
            local_candidates_path = temp_path / "local_candidates.tsv"
            local_manifest_path = temp_path / "manifest.tsv"
            output_dir = temp_path / "authority"

            write_tsv(
                families_path,
                [
                    {
                        "family_id": "fam-raw-u-tin-htway-first-burmese-royal-inscription",
                        "family_label": "u tin htway first burmese royal inscription",
                        "family_type": "article",
                        "member_count": "1",
                        "occurrence_count": "2",
                        "sample_raw_references": "U Tin Htway, First Burmese Royal Inscription",
                        "likely_contains_translation": "unknown",
                        "review_status": "needs_human_review",
                        "notes": "",
                    }
                ],
                FAMILY_FIELDS,
            )
            write_tsv(
                members_path,
                [
                    {
                        "family_id": "fam-raw-u-tin-htway-first-burmese-royal-inscription",
                        "raw_reference_string": "U Tin Htway, First Burmese Royal Inscription",
                        "occurrence_count": "2",
                        "example_record_ids": "obi-1",
                        "notes": "",
                    }
                ],
                MEMBER_FIELDS,
            )
            write_tsv(
                candidates_path,
                [
                    {
                        "work_candidate_id": "wc-tin-htway",
                        "family_id": "fam-raw-u-tin-htway-first-burmese-royal-inscription",
                        "provisional_short_label": "First Burmese Royal Inscription",
                        "author_original": "U Tin Htway",
                        "author_normalized": "u tin htway",
                        "year": "",
                        "title_original": "First Burmese Royal Inscription",
                        "title_normalized": "first burmese royal inscription",
                        "publication_details": "",
                        "language": "",
                        "script": "",
                        "translation_relevance": "unknown",
                        "evidence_raw_references": "U Tin Htway, First Burmese Royal Inscription",
                        "review_status": "needs_human_review",
                        "notes": "",
                    }
                ],
                CANDIDATE_FIELDS,
            )
            write_tsv(seeds_path, [], SEED_FIELDS)
            write_tsv(
                local_candidates_path,
                [
                    {
                        "candidate_id": "tinhtway1974",
                        "search_term": "u tin htswe",
                        "name": "Tin Htway 1974 Oldest Burmese Inscription.pdf",
                        "original_path": "OBI_LIBRARY_ROOT:Thematic/Pyu/Tin Htway 1974 Oldest Burmese Inscription.pdf",
                        "file_type": "pdf",
                        "file_size": "100",
                        "sha256": "abc",
                        "probable_work_label": "Oldest Burmese Inscription",
                        "probable_author": "Tin Htway",
                        "probable_year": "1974",
                        "match_confidence": "high",
                        "copy_status": "copied",
                        "copied_path": "data/local/bibliography_sources/tinhtway1974/Tin Htway 1974 Oldest Burmese Inscription.pdf",
                        "notes": "",
                    }
                ],
                LOCAL_CANDIDATE_FIELDS,
            )
            acronym_candidates_path, acronym_report_path = self.write_fixture_acronym_files(temp_path)
            write_tsv(
                local_manifest_path,
                [
                    {
                        "source_file_id": "tinhtway1974",
                        "original_path": "OBI_LIBRARY_ROOT:Thematic/Pyu/Tin Htway 1974 Oldest Burmese Inscription.pdf",
                        "copied_path": "data/local/bibliography_sources/tinhtway1974/Tin Htway 1974 Oldest Burmese Inscription.pdf",
                        "file_name": "Tin Htway 1974 Oldest Burmese Inscription.pdf",
                        "file_type": "pdf",
                        "file_size": "100",
                        "sha256": "abc",
                        "source_folder_hint": "Pyu",
                        "copy_date": "2024-01-01T00:00:00Z",
                        "copy_status": "copied",
                        "notes": "",
                    }
                ],
                LOCAL_MANIFEST_FIELDS,
            )

            report = build_authority(
                reference_families_path=families_path,
                reference_members_path=members_path,
                work_candidates_path=candidates_path,
                seed_path=seeds_path,
                output_dir=output_dir,
                local_candidates_path=local_candidates_path,
                local_manifest_path=local_manifest_path,
                acronym_candidates_path=acronym_candidates_path,
                acronym_report_path=acronym_report_path,
            )

            authority = (output_dir / "bibtex_authority.tsv").read_text(encoding="utf-8")
            self.assertIn("confirmed_local_source", authority)
            self.assertIn("Oldest Burmese Inscription", authority)
            self.assertEqual(report["confirmed_local_source_count"], 1)

    def test_validate_bibtex_authority_rejects_unsorted_high_frequency_output(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            families_path, members_path, candidates_path, seeds_path, external_entries_path = self.write_fixture_tables(temp_path)
            acronym_candidates_path, acronym_report_path = self.write_fixture_acronym_files(temp_path)
            manual_seed_path = self.write_manual_seed_file(temp_path)
            output_dir = temp_path / "authority"
            build_authority(
                reference_families_path=families_path,
                reference_members_path=members_path,
                work_candidates_path=candidates_path,
                seed_path=seeds_path,
                external_entries_path=external_entries_path,
                output_dir=output_dir,
                frasch_references_path=temp_path / "missing_frasch.tsv",
                local_candidates_path=temp_path / "missing_local_candidates.tsv",
                local_manifest_path=temp_path / "missing_local_manifest.tsv",
                acronym_candidates_path=acronym_candidates_path,
                acronym_report_path=acronym_report_path,
            )
            high_frequency_path = output_dir / "high_frequency_unresolved.tsv"
            high_frequency_path.write_text(
                "family_id\tfamily_label\tfamily_type\tmember_count\toccurrence_count\tsample_raw_references\tcurrent_bibtex_key\tcurrent_status\tsuggested_local_search_terms\tnotes\n"
                "fam-a\tA\tunclear\t1\t1\tA 1\tkey-a\tmachine_stub\tA\t\n"
                "fam-b\tB\tunclear\t1\t10\tB 1\tkey-b\tmachine_stub\tB\t\n",
                encoding="utf-8",
            )
            result = validate_bibtex_authority(
                authority_bib_path=output_dir / "bibliography_authority.bib",
                candidates_bib_path=output_dir / "bibliography_candidates.bib",
                authority_tsv_path=output_dir / "bibtex_authority.tsv",
                crosswalk_path=output_dir / "raw_reference_to_bibtex.tsv",
                families_path=families_path,
                external_entries_path=external_entries_path,
                seed_path=seeds_path,
                high_frequency_path=high_frequency_path,
                acronym_status_path=output_dir / "acronym_resolution_status.tsv",
                acronym_candidates_path=acronym_candidates_path,
                acronym_report_path=acronym_report_path,
            )
            self.assertFalse(result["ok"])
            self.assertTrue(any("not sorted" in error for error in result["errors"]))

    def test_validate_bibtex_authority_passes_for_valid_fixture(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            families_path, members_path, candidates_path, seeds_path, external_entries_path = self.write_fixture_tables(temp_path)
            acronym_candidates_path, acronym_report_path = self.write_fixture_acronym_files(temp_path)
            manual_seed_path = self.write_manual_seed_file(temp_path)
            output_dir = temp_path / "authority"
            build_authority(
                reference_families_path=families_path,
                reference_members_path=members_path,
                work_candidates_path=candidates_path,
                seed_path=seeds_path,
                external_entries_path=external_entries_path,
                output_dir=output_dir,
                frasch_references_path=temp_path / "missing_frasch.tsv",
                local_candidates_path=temp_path / "missing_local_candidates.tsv",
                local_manifest_path=temp_path / "missing_local_manifest.tsv",
                acronym_candidates_path=acronym_candidates_path,
                acronym_report_path=acronym_report_path,
                manual_acronym_seeds_path=manual_seed_path,
            )
            result = validate_bibtex_authority(
                authority_bib_path=output_dir / "bibliography_authority.bib",
                candidates_bib_path=output_dir / "bibliography_candidates.bib",
                authority_tsv_path=output_dir / "bibtex_authority.tsv",
                crosswalk_path=output_dir / "raw_reference_to_bibtex.tsv",
                families_path=families_path,
                external_entries_path=external_entries_path,
                evidence_path=output_dir / "bibtex_authority_evidence.tsv",
                acronym_status_path=output_dir / "acronym_resolution_status.tsv",
                acronym_candidates_path=acronym_candidates_path,
                acronym_report_path=acronym_report_path,
                manual_acronym_seeds_path=manual_seed_path,
                source_family_path=output_dir / "source_family_authority.tsv",
                manual_review_packet_path=output_dir / "acronym_manual_review_packet.tsv",
                remaining_acronym_worklist_path=output_dir / "remaining_acronym_worklist.tsv",
                remaining_acronym_evidence_path=output_dir / "remaining_acronym_evidence.tsv",
                source_work_locator_systems_path=output_dir / "source_work_locator_systems.tsv",
                source_work_authority_path=output_dir / "source_work_authority.tsv",
                source_work_authority_audit_path=output_dir / "source_work_authority_audit.tsv",
                source_work_to_bibtex_reconciliation_path=output_dir / "source_work_to_bibtex_reconciliation.tsv",
                bibtex_field_quality_audit_path=output_dir / "bibtex_field_quality_audit.tsv",
                authority_key_normalization_path=output_dir / "authority_key_normalization.tsv",
                raw_reference_crosswalk_audit_path=output_dir / "raw_reference_crosswalk_audit.tsv",
                candidate_stub_review_path=output_dir / "candidate_stub_review.tsv",
            )
            self.assertTrue(result["ok"], result["errors"])

    def test_build_authority_emits_source_work_qc_tables(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            families_path, members_path, candidates_path, seeds_path, external_entries_path = self.write_fixture_tables(temp_path)
            acronym_candidates_path, acronym_report_path = self.write_fixture_acronym_files(temp_path)
            output_dir = temp_path / "authority"
            build_authority(
                reference_families_path=families_path,
                reference_members_path=members_path,
                work_candidates_path=candidates_path,
                seed_path=seeds_path,
                external_entries_path=external_entries_path,
                output_dir=output_dir,
                frasch_references_path=temp_path / "missing_frasch.tsv",
                local_candidates_path=temp_path / "missing_local_candidates.tsv",
                local_manifest_path=temp_path / "missing_local_manifest.tsv",
                acronym_candidates_path=acronym_candidates_path,
                acronym_report_path=acronym_report_path,
            )

            source_work_rows = {row["source_work_key"]: row for row in read_tsv(output_dir / "source_work_authority.tsv")}
            reconciliation_rows = {row["source_work_key"]: row for row in read_tsv(output_dir / "source_work_to_bibtex_reconciliation.tsv")}
            normalization_rows = read_tsv(output_dir / "authority_key_normalization.tsv")
            report = json.loads((output_dir / "bibtex_authority_report.json").read_text(encoding="utf-8"))

            self.assertIn("oldBurmeseInscriptions", source_work_rows)
            self.assertTrue(source_work_rows["oldBurmeseInscriptions"]["authority_level"])
            self.assertIn("duroiselle1921list", reconciliation_rows)
            self.assertEqual(reconciliation_rows["duroiselle1921list"]["bibtex_status"], "present")
            self.assertTrue(any(row["old_key"] == "obiCorpusSource" and row["new_key"] == "oldBurmeseInscriptions" for row in normalization_rows))
            self.assertIn("source_work_authority_audit_count", report)
            self.assertIn("bibtex_field_quality_issue_count", report)
            self.assertIn("authority_key_normalization_count", report)

    def test_validate_bibtex_authority_rejects_missing_source_work_qc_artifacts(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            families_path, members_path, candidates_path, seeds_path, external_entries_path = self.write_fixture_tables(temp_path)
            acronym_candidates_path, acronym_report_path = self.write_fixture_acronym_files(temp_path)
            output_dir = temp_path / "authority"
            build_authority(
                reference_families_path=families_path,
                reference_members_path=members_path,
                work_candidates_path=candidates_path,
                seed_path=seeds_path,
                external_entries_path=external_entries_path,
                output_dir=output_dir,
                frasch_references_path=temp_path / "missing_frasch.tsv",
                local_candidates_path=temp_path / "missing_local_candidates.tsv",
                local_manifest_path=temp_path / "missing_local_manifest.tsv",
                acronym_candidates_path=acronym_candidates_path,
                acronym_report_path=acronym_report_path,
            )
            (output_dir / "source_work_authority_audit.tsv").unlink()

            result = validate_bibtex_authority(
                authority_bib_path=output_dir / "bibliography_authority.bib",
                candidates_bib_path=output_dir / "bibliography_candidates.bib",
                authority_tsv_path=output_dir / "bibtex_authority.tsv",
                crosswalk_path=output_dir / "raw_reference_to_bibtex.tsv",
                families_path=families_path,
                external_entries_path=external_entries_path,
                evidence_path=output_dir / "bibtex_authority_evidence.tsv",
                acronym_status_path=output_dir / "acronym_resolution_status.tsv",
                acronym_candidates_path=acronym_candidates_path,
                acronym_report_path=acronym_report_path,
                source_family_path=output_dir / "source_family_authority.tsv",
                manual_review_packet_path=output_dir / "acronym_manual_review_packet.tsv",
                remaining_acronym_worklist_path=output_dir / "remaining_acronym_worklist.tsv",
                remaining_acronym_evidence_path=output_dir / "remaining_acronym_evidence.tsv",
                source_work_locator_systems_path=output_dir / "source_work_locator_systems.tsv",
                source_work_authority_path=output_dir / "source_work_authority.tsv",
                source_work_authority_audit_path=output_dir / "source_work_authority_audit.tsv",
                source_work_to_bibtex_reconciliation_path=output_dir / "source_work_to_bibtex_reconciliation.tsv",
                bibtex_field_quality_audit_path=output_dir / "bibtex_field_quality_audit.tsv",
                authority_key_normalization_path=output_dir / "authority_key_normalization.tsv",
                raw_reference_crosswalk_audit_path=output_dir / "raw_reference_crosswalk_audit.tsv",
                candidate_stub_review_path=output_dir / "candidate_stub_review.tsv",
            )
            self.assertFalse(result["ok"])
            self.assertTrue(any("source_work_authority_audit.tsv is missing" in error for error in result["errors"]))

    def test_validate_bibtex_authority_requires_strong_acronym_evidence(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            families_path, members_path, candidates_path, seeds_path, external_entries_path = self.write_fixture_tables(temp_path)
            acronym_candidates_path, acronym_report_path = self.write_fixture_acronym_files(temp_path)
            output_dir = temp_path / "authority"
            build_authority(
                reference_families_path=families_path,
                reference_members_path=members_path,
                work_candidates_path=candidates_path,
                seed_path=seeds_path,
                external_entries_path=external_entries_path,
                output_dir=output_dir,
                frasch_references_path=temp_path / "missing_frasch.tsv",
                local_candidates_path=temp_path / "missing_local_candidates.tsv",
                local_manifest_path=temp_path / "missing_local_manifest.tsv",
                acronym_candidates_path=acronym_candidates_path,
                acronym_report_path=acronym_report_path,
            )
            acronym_status_path = output_dir / "acronym_resolution_status.tsv"
            rows = read_tsv(acronym_status_path)
            ppa_row = next(row for row in rows if row["acronym"] == "PPA")
            ppa_row["resolution_status"] = "confirmed_expansion"
            ppa_row["best_evidence_id"] = "doc:pl"
            ppa_row["best_evidence_quote"] = "Pl. II 198"
            write_tsv(acronym_status_path, rows, ACRONYM_STATUS_FIELDS)
            result = validate_bibtex_authority(
                authority_bib_path=output_dir / "bibliography_authority.bib",
                candidates_bib_path=output_dir / "bibliography_candidates.bib",
                authority_tsv_path=output_dir / "bibtex_authority.tsv",
                crosswalk_path=output_dir / "raw_reference_to_bibtex.tsv",
                families_path=families_path,
                external_entries_path=external_entries_path,
                source_family_path=output_dir / "source_family_authority.tsv",
                evidence_path=output_dir / "bibtex_authority_evidence.tsv",
                report_path=output_dir / "bibtex_authority_report.json",
                acronym_status_path=acronym_status_path,
                acronym_candidates_path=acronym_candidates_path,
                acronym_report_path=acronym_report_path,
            )
            self.assertFalse(result["ok"])
            self.assertTrue(any("requires strong evidence" in error for error in result["errors"]))

    def test_validate_bibtex_authority_rejects_long_acronym_quote(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            families_path, members_path, candidates_path, seeds_path, external_entries_path = self.write_fixture_tables(temp_path)
            acronym_candidates_path, acronym_report_path = self.write_fixture_acronym_files(temp_path)
            output_dir = temp_path / "authority"
            build_authority(
                reference_families_path=families_path,
                reference_members_path=members_path,
                work_candidates_path=candidates_path,
                seed_path=seeds_path,
                external_entries_path=external_entries_path,
                output_dir=output_dir,
                frasch_references_path=temp_path / "missing_frasch.tsv",
                local_candidates_path=temp_path / "missing_local_candidates.tsv",
                local_manifest_path=temp_path / "missing_local_manifest.tsv",
                acronym_candidates_path=acronym_candidates_path,
                acronym_report_path=acronym_report_path,
            )
            acronym_status_path = output_dir / "acronym_resolution_status.tsv"
            rows = read_tsv(acronym_status_path)
            sip_row = next(row for row in rows if row["acronym"] == "SIP")
            sip_row["best_evidence_quote"] = "x" * 240
            write_tsv(acronym_status_path, rows, ACRONYM_STATUS_FIELDS)
            result = validate_bibtex_authority(
                authority_bib_path=output_dir / "bibliography_authority.bib",
                candidates_bib_path=output_dir / "bibliography_candidates.bib",
                authority_tsv_path=output_dir / "bibtex_authority.tsv",
                crosswalk_path=output_dir / "raw_reference_to_bibtex.tsv",
                families_path=families_path,
                external_entries_path=external_entries_path,
                source_family_path=output_dir / "source_family_authority.tsv",
                evidence_path=output_dir / "bibtex_authority_evidence.tsv",
                report_path=output_dir / "bibtex_authority_report.json",
                acronym_status_path=acronym_status_path,
                acronym_candidates_path=acronym_candidates_path,
                acronym_report_path=acronym_report_path,
            )
            self.assertFalse(result["ok"])
            self.assertTrue(any("best_evidence_quote exceeds" in error for error in result["errors"]))

    def test_harvest_local_bibliography_sources_collapses_duplicates_by_sha(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            author_root = temp_path / "Authors alphabetical"
            frasch_dir = author_root / "Frasch, Tilmans"
            library_dir = temp_path / "Library" / "Frasch"
            frasch_dir.mkdir(parents=True)
            library_dir.mkdir(parents=True)
            duplicate_text = "same content"
            (frasch_dir / "Bagan Epig Database.doc").write_text(duplicate_text, encoding="utf-8")
            (library_dir / "Bagan Epig Database copy.doc").write_text(duplicate_text, encoding="utf-8")
            downloads = temp_path / "Downloads"
            downloads.mkdir()

            previous = {key: os.environ.get(key) for key in ("OBI_AUTHOR_ALPHA_ROOT", "OBI_LIBRARY_ROOT", "OBI_LOCAL_BIB_ROOT")}
            os.environ["OBI_AUTHOR_ALPHA_ROOT"] = str(author_root)
            os.environ["OBI_LIBRARY_ROOT"] = str(temp_path / "Library")
            os.environ["OBI_LOCAL_BIB_ROOT"] = str(downloads)
            try:
                output_dir = temp_path / "output"
                report = run_harvest("frasch", output_dir)
            finally:
                for key, value in previous.items():
                    if value is None:
                        os.environ.pop(key, None)
                    else:
                        os.environ[key] = value

            self.assertEqual(report["raw_candidate_count"], 2)
            self.assertEqual(report["unique_file_count"], 1)
            self.assertEqual(report["duplicate_file_count"], 1)
            manifest = (output_dir / "local_file_manifest.tsv").read_text(encoding="utf-8")
            self.assertIn("Bagan Epig Database copy.doc", manifest)

    def test_build_authority_writes_evidence_rows_and_clips_bibtex_evidence(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            families_path, members_path, candidates_path, seeds_path, external_entries_path = self.write_fixture_tables(temp_path)
            frasch_path = temp_path / "frasch.tsv"
            long_reference = "Harvey, History of Burma, " + ("very long evidence " * 20)
            write_tsv(
                frasch_path,
                [
                    {
                        "frasch_ref_id": "frasch-ref-1",
                        "raw_reference": long_reference,
                        "author": "Harvey",
                        "editor": "",
                        "year": "1925",
                        "title": "History of Burma",
                        "publication": "",
                        "journal": "",
                        "volume": "",
                        "number": "",
                        "pages": "",
                        "publisher": "",
                        "place": "",
                        "language": "latin",
                        "script": "Latn",
                        "confidence": "high",
                        "detected_entry_type": "bibliographic_reference",
                        "looks_like_bibliographic_reference": "true",
                        "looks_like_catalogue_note": "false",
                        "looks_like_body_text": "false",
                        "has_author_signal": "true",
                        "has_year_signal": "true",
                        "has_title_signal": "true",
                        "has_publication_signal": "true",
                        "length": str(len(long_reference)),
                        "recommended_action": "use_for_bibliography",
                        "source_location_hint": "fixture",
                        "extraction_source_file": "frasch.doc",
                        "notes": "",
                    }
                ],
                [
                    "frasch_ref_id",
                    "raw_reference",
                    "author",
                    "editor",
                    "year",
                    "title",
                    "publication",
                    "journal",
                    "volume",
                    "number",
                    "pages",
                    "publisher",
                    "place",
                    "language",
                    "script",
                    "confidence",
                    "detected_entry_type",
                    "looks_like_bibliographic_reference",
                    "looks_like_catalogue_note",
                    "looks_like_body_text",
                    "has_author_signal",
                    "has_year_signal",
                    "has_title_signal",
                    "has_publication_signal",
                    "length",
                    "recommended_action",
                    "source_location_hint",
                    "extraction_source_file",
                    "notes",
                ],
            )
            output_dir = temp_path / "authority"
            acronym_candidates_path, acronym_report_path = self.write_fixture_acronym_files(temp_path)
            build_authority(
                reference_families_path=families_path,
                reference_members_path=members_path,
                work_candidates_path=candidates_path,
                seed_path=seeds_path,
                external_entries_path=external_entries_path,
                frasch_references_path=frasch_path,
                output_dir=output_dir,
                local_candidates_path=temp_path / "missing_local_candidates.tsv",
                local_manifest_path=temp_path / "missing_local_manifest.tsv",
                acronym_candidates_path=acronym_candidates_path,
                acronym_report_path=acronym_report_path,
            )
            evidence_tsv = (output_dir / "bibtex_authority_evidence.tsv").read_text(encoding="utf-8")
            authority_bib = (output_dir / "bibliography_authority.bib").read_text(encoding="utf-8")
            self.assertIn("frasch-ref-1", evidence_tsv)
            self.assertNotIn(long_reference, authority_bib)
            self.assertIn("harvey1925historyBurma", authority_bib)

    def test_validate_bibtex_authority_rejects_long_matched_local_reference(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            families_path = temp_path / "reference_families.tsv"
            crosswalk_path = temp_path / "raw_reference_to_bibtex.tsv"
            authority_tsv_path = temp_path / "bibtex_authority.tsv"
            authority_bib_path = temp_path / "bibliography_authority.bib"
            candidate_bib_path = temp_path / "bibliography_candidates.bib"
            external_entries_path = temp_path / "external.tsv"
            long_reference = "x" * 200

            write_tsv(families_path, [{"family_id": "fam-1", "family_label": "OBI", "family_type": "source_catalogue", "member_count": "1", "occurrence_count": "1", "sample_raw_references": "OBI 1", "likely_contains_translation": "no", "review_status": "unreviewed", "notes": ""}], FAMILY_FIELDS)
            write_tsv(
                authority_tsv_path,
                [
                    {
                        "bibtex_key": "obiCorpusSource",
                        "entry_type": "misc",
                        "authority_status": "confirmed_local_source",
                        "source_of_authority": "local_burma_folder",
                        "matched_external_key": "",
                        "matched_local_source_id": "local-1",
                        "matched_local_source_file": "source.pdf",
                        "matched_local_reference": long_reference,
                        "match_confidence": "high",
                        "match_reason": "fixture",
                        "evidence_id": "local-1",
                        "short_evidence_note": "fixture",
                        "human_review_flag": "false",
                        "resolution_status": "source_family_resolved",
                        "resolution_level": "internal_reference",
                        "source_family_id": "sf-obi",
                        "family_id": "fam-1",
                        "family_label": "OBI",
                        "family_type": "source_catalogue",
                        "author": "",
                        "editor": "",
                        "year": "",
                        "title": "OBI",
                        "shorttitle": "OBI",
                        "journal": "",
                        "booktitle": "",
                        "publisher": "",
                        "address": "",
                        "volume": "",
                        "number": "",
                        "pages": "",
                        "doi": "",
                        "url": "",
                        "isbn": "",
                        "language": "",
                        "script": "",
                        "translation_relevance": "unknown",
                        "review_status": "reviewed_confirmed",
                        "evidence": "fixture",
                        "notes": "",
                    }
                ],
                AUTHORITY_FIELDS,
            )
            authority_bib_path.write_text(
                "@misc{obiCorpusSource,\n"
                f"  matchedlocalreference = {{{long_reference}}}\n"
                "}\n",
                encoding="utf-8",
            )
            candidate_bib_path.write_text("", encoding="utf-8")
            write_tsv(
                crosswalk_path,
                [
                    {
                        "raw_reference_string": "OBI 1",
                        "family_id": "fam-1",
                        "source_family_id": "sf-obi",
                        "work_candidate_id": "",
                        "bibtex_key": "obiCorpusSource",
                        "locator": "",
                        "locator_type": "number",
                        "resolution_status": "source_family_resolved",
                        "resolution_level": "internal_reference",
                        "match_type": "fixture",
                        "match_confidence": "high",
                        "evidence": "fixture",
                        "needs_human_review": "false",
                        "notes": "",
                    }
                ],
                CROSSWALK_FIELDS,
            )
            write_tsv(
                temp_path / "source_family_authority.tsv",
                [
                    {
                        "source_family_id": "sf-obi",
                        "abbreviation": "OBI",
                        "family_id": "fam-1",
                        "authority_key": "obiCorpusSource",
                        "source_family_type": "corpus_internal",
                        "resolution_status": "source_family_resolved",
                        "resolution_level": "internal_reference",
                        "canonical_label": "OBI",
                        "expanded_label": "Old Burmese Inscriptions",
                        "related_bibtex_key": "obiCorpusSource",
                        "locator_pattern": "number",
                        "example_raw_references": "OBI 1",
                        "evidence_id": "local-1",
                        "evidence_source": "local_burma_folder",
                        "confidence": "high",
                        "needs_human_review": "false",
                        "notes": "",
                    }
                ],
                SOURCE_FAMILY_FIELDS,
            )
            write_tsv(
                temp_path / "bibtex_authority_evidence.tsv",
                [
                    {
                        "source_family_id": "sf-obi",
                        "bibtex_key": "obiCorpusSource",
                        "evidence_id": "local-1",
                        "evidence_type": "local_burma_folder",
                        "source_file_id": "local-1",
                        "source_file_label": "source.pdf",
                        "source_ref_id": "local-1",
                        "short_evidence": "fixture",
                        "full_evidence_hash": "abc",
                        "confidence": "high",
                        "notes": "",
                    }
                ],
                ["source_family_id", "bibtex_key", "evidence_id", "evidence_type", "source_file_id", "source_file_label", "source_ref_id", "short_evidence", "full_evidence_hash", "confidence", "notes"],
            )
            write_tsv(external_entries_path, [], ["bibtex_key"])
            result = validate_bibtex_authority(
                authority_bib_path=authority_bib_path,
                candidates_bib_path=candidate_bib_path,
                authority_tsv_path=authority_tsv_path,
                crosswalk_path=crosswalk_path,
                families_path=families_path,
                external_entries_path=external_entries_path,
                source_family_path=temp_path / "source_family_authority.tsv",
                evidence_path=temp_path / "bibtex_authority_evidence.tsv",
            )
            self.assertFalse(result["ok"])
            self.assertTrue(any("matchedlocalreference" in error or "matched_local_reference" in error for error in result["errors"]))

    def test_build_authority_resolves_b2_family_to_shared_seed_key(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            families_path = temp_path / "families.tsv"
            members_path = temp_path / "members.tsv"
            candidates_path = temp_path / "candidates.tsv"
            seeds_path = temp_path / "seeds.tsv"
            output_dir = temp_path / "authority"

            write_tsv(
                families_path,
                [
                    {"family_id": "fam-raw-b", "family_label": "B", "family_type": "source_catalogue", "member_count": "1", "occurrence_count": "5", "sample_raw_references": "B 1", "likely_contains_translation": "no", "review_status": "needs_human_review", "notes": ""},
                    {"family_id": "fam-raw-b-2", "family_label": "B 2, p. 815", "family_type": "source_catalogue", "member_count": "1", "occurrence_count": "4", "sample_raw_references": "B 2, p. 815", "likely_contains_translation": "no", "review_status": "needs_human_review", "notes": ""},
                ],
                FAMILY_FIELDS,
            )
            write_tsv(
                members_path,
                [
                    {"family_id": "fam-raw-b", "raw_reference_string": "B 1", "occurrence_count": "5", "example_record_ids": "obi-1", "notes": ""},
                    {"family_id": "fam-raw-b-2", "raw_reference_string": "B 2, p. 815", "occurrence_count": "4", "example_record_ids": "obi-2", "notes": ""},
                ],
                MEMBER_FIELDS,
            )
            write_tsv(candidates_path, [], CANDIDATE_FIELDS)
            write_tsv(
                seeds_path,
                [
                    {
                        "abbreviation": "B",
                        "family_id": "fam-raw-b",
                        "family_type": "source_catalogue",
                        "provisional_label": "Bagan Epigraphic Database, Part B",
                        "probable_bibtex_key": "fraschBaganEpigraphicDatabasePartB",
                        "source_type": "source_catalogue",
                        "evidence_source_file": "Bagan Epig Database.doc",
                        "evidence_ref_id": "",
                        "evidence_quote_short": "B",
                        "confidence": "medium",
                        "needs_human_review": "true",
                        "notes": "",
                    }
                ],
                SEED_FIELDS,
            )
            acronym_candidates_path, acronym_report_path = self.write_fixture_acronym_files(temp_path)
            build_authority(
                reference_families_path=families_path,
                reference_members_path=members_path,
                work_candidates_path=candidates_path,
                seed_path=seeds_path,
                output_dir=output_dir,
                frasch_references_path=temp_path / "missing_frasch.tsv",
                local_candidates_path=temp_path / "missing_local_candidates.tsv",
                local_manifest_path=temp_path / "missing_local_manifest.tsv",
                acronym_candidates_path=acronym_candidates_path,
                acronym_report_path=acronym_report_path,
            )
            crosswalk_rows = read_tsv(output_dir / "raw_reference_to_bibtex.tsv")
            resolution_rows = read_tsv(output_dir / "high_frequency_resolution_plan.tsv")
            unresolved = (output_dir / "high_frequency_unresolved.tsv").read_text(encoding="utf-8")
            b2_row = next(row for row in crosswalk_rows if row["family_id"] == "fam-raw-b-2")
            self.assertEqual(b2_row["source_family_id"], "sf-b")
            self.assertEqual(b2_row["resolution_status"], "alias_resolved")
            self.assertEqual(b2_row["locator_type"], "volume_page")
            plan_row = next(row for row in resolution_rows if row["family_id"] == "fam-raw-b-2")
            self.assertEqual(plan_row["resolution_status"], "alias_resolved")
            self.assertEqual(plan_row["authority_key"], "fraschBaganEpigraphicDatabasePartB")
            self.assertNotIn("fam-raw-b-2", unresolved)

    def test_build_authority_models_source_family_locators_without_machine_stubs(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            families_path = temp_path / "families.tsv"
            members_path = temp_path / "members.tsv"
            candidates_path = temp_path / "candidates.tsv"
            seeds_path = temp_path / "seeds.tsv"
            output_dir = temp_path / "authority"

            write_tsv(
                families_path,
                [
                    {"family_id": "fam-rdasb-publication", "family_label": "RDASB", "family_type": "publication", "member_count": "1", "occurrence_count": "9", "sample_raw_references": "RDASB 1971", "likely_contains_translation": "no", "review_status": "needs_human_review", "notes": ""},
                    {"family_id": "fam-plate-references", "family_label": "Pl.", "family_type": "internal_reference", "member_count": "1", "occurrence_count": "8", "sample_raw_references": "Pl. II 198", "likely_contains_translation": "no", "review_status": "needs_human_review", "notes": ""},
                    {"family_id": "fam-list-catalogue", "family_label": "List", "family_type": "source_catalogue", "member_count": "1", "occurrence_count": "7", "sample_raw_references": "List 90", "likely_contains_translation": "no", "review_status": "needs_human_review", "notes": ""},
                    {"family_id": "fam-raw-ub", "family_label": "UB", "family_type": "source_catalogue", "member_count": "1", "occurrence_count": "6", "sample_raw_references": "UB 1, p. 297", "likely_contains_translation": "no", "review_status": "needs_human_review", "notes": ""},
                    {"family_id": "fam-raw-mp", "family_label": "MP", "family_type": "source_catalogue", "member_count": "1", "occurrence_count": "5", "sample_raw_references": "MP 1, p. 81", "likely_contains_translation": "no", "review_status": "needs_human_review", "notes": ""},
                    {"family_id": "fam-ppa-catalogue", "family_label": "PPA", "family_type": "source_catalogue", "member_count": "1", "occurrence_count": "4", "sample_raw_references": "PPA, p. 55", "likely_contains_translation": "no", "review_status": "needs_human_review", "notes": ""},
                    {"family_id": "fam-raw-a", "family_label": "A", "family_type": "source_catalogue", "member_count": "1", "occurrence_count": "3", "sample_raw_references": "A, p. 79", "likely_contains_translation": "no", "review_status": "needs_human_review", "notes": ""},
                    {"family_id": "fam-raw-b", "family_label": "B", "family_type": "source_catalogue", "member_count": "1", "occurrence_count": "2", "sample_raw_references": "B 2, p. 815", "likely_contains_translation": "no", "review_status": "needs_human_review", "notes": ""},
                ],
                FAMILY_FIELDS,
            )
            write_tsv(
                members_path,
                [
                    {"family_id": "fam-rdasb-publication", "raw_reference_string": "RDASB 1971", "occurrence_count": "9", "example_record_ids": "obi-1", "notes": ""},
                    {"family_id": "fam-plate-references", "raw_reference_string": "Pl. II 198", "occurrence_count": "8", "example_record_ids": "obi-2", "notes": ""},
                    {"family_id": "fam-list-catalogue", "raw_reference_string": "List 90", "occurrence_count": "7", "example_record_ids": "obi-3", "notes": ""},
                    {"family_id": "fam-raw-ub", "raw_reference_string": "UB 1, p. 297", "occurrence_count": "6", "example_record_ids": "obi-4", "notes": ""},
                    {"family_id": "fam-raw-mp", "raw_reference_string": "MP 1, p. 81", "occurrence_count": "5", "example_record_ids": "obi-5", "notes": ""},
                    {"family_id": "fam-ppa-catalogue", "raw_reference_string": "PPA, p. 55", "occurrence_count": "4", "example_record_ids": "obi-6", "notes": ""},
                    {"family_id": "fam-raw-a", "raw_reference_string": "A, p. 79", "occurrence_count": "3", "example_record_ids": "obi-7", "notes": ""},
                    {"family_id": "fam-raw-b", "raw_reference_string": "B 2, p. 815", "occurrence_count": "2", "example_record_ids": "obi-8", "notes": ""},
                ],
                MEMBER_FIELDS,
            )
            write_tsv(candidates_path, [], CANDIDATE_FIELDS)
            write_tsv(
                seeds_path,
                [
                    {"abbreviation": "List", "family_id": "fam-list-catalogue", "provisional_label": "List of Inscriptions Found in Burma", "probable_bibtex_key": "duroiselle1921list", "source_type": "source_catalogue", "confidence": "high", "evidence": "List 90", "needs_human_review": "false", "notes": ""},
                    {"abbreviation": "RDASB", "family_id": "fam-rdasb-publication", "provisional_label": "Report of the Director, Archaeological Survey of Burma", "probable_bibtex_key": "reportDirectorArchaeologicalSurveyBurma", "source_type": "publication", "confidence": "high", "evidence": "RDASB 1971", "needs_human_review": "false", "notes": ""},
                    {"abbreviation": "UB", "family_id": "fam-raw-ub", "provisional_label": "UB source family", "probable_bibtex_key": "ubSourceFamily", "source_type": "source_catalogue", "confidence": "medium", "evidence": "UB 1, p. 297", "needs_human_review": "true", "notes": ""},
                    {"abbreviation": "MP", "family_id": "fam-raw-mp", "provisional_label": "MP source family", "probable_bibtex_key": "mpSourceFamily", "source_type": "source_catalogue", "confidence": "medium", "evidence": "MP 1, p. 81", "needs_human_review": "true", "notes": ""},
                    {"abbreviation": "PPA", "family_id": "fam-ppa-catalogue", "provisional_label": "PPA source family", "probable_bibtex_key": "ppaCatalogueFamily", "source_type": "source_catalogue", "confidence": "medium", "evidence": "PPA, p. 55", "needs_human_review": "true", "notes": ""},
                    {"abbreviation": "A", "family_id": "fam-raw-a", "provisional_label": "Bagan Epigraphic Database, Part A", "probable_bibtex_key": "fraschBaganEpigraphicDatabasePartA", "source_type": "source_catalogue", "confidence": "medium", "evidence": "A, p. 79", "needs_human_review": "true", "notes": ""},
                    {"abbreviation": "B", "family_id": "fam-raw-b", "provisional_label": "Bagan Epigraphic Database, Part B", "probable_bibtex_key": "fraschBaganEpigraphicDatabasePartB", "source_type": "source_catalogue", "confidence": "medium", "evidence": "B 2, p. 815", "needs_human_review": "true", "notes": ""},
                ],
                SEED_FIELDS,
            )
            acronym_candidates_path, acronym_report_path = self.write_fixture_acronym_files(temp_path)

            report = build_authority(
                reference_families_path=families_path,
                reference_members_path=members_path,
                work_candidates_path=candidates_path,
                seed_path=seeds_path,
                output_dir=output_dir,
                frasch_references_path=temp_path / "missing_frasch.tsv",
                local_candidates_path=temp_path / "missing_local_candidates.tsv",
                local_manifest_path=temp_path / "missing_local_manifest.tsv",
                acronym_candidates_path=acronym_candidates_path,
                acronym_report_path=acronym_report_path,
            )

            crosswalk_rows = read_tsv(output_dir / "raw_reference_to_bibtex.tsv")
            source_family_rows = {row["source_family_id"]: row for row in read_tsv(output_dir / "source_family_authority.tsv")}
            acronym_rows = {row["acronym"]: row for row in read_tsv(output_dir / "acronym_resolution_status.tsv")}
            locator_rows = read_tsv(output_dir / "source_work_locator_systems.tsv")
            remaining_worklist_rows = read_tsv(output_dir / "remaining_acronym_worklist.tsv")
            final_sprint_rows = {row["acronym"]: row for row in read_tsv(output_dir / "final_acronym_resolution_sprint.tsv")}
            unresolved_dossier_rows = read_tsv(output_dir / "unresolved_acronym_dossier.tsv")
            candidate_bib = (output_dir / "bibliography_candidates.bib").read_text(encoding="utf-8")

            by_family = {row["family_id"]: row for row in crosswalk_rows}
            self.assertEqual(by_family["fam-rdasb-publication"]["resolution_status"], "series_level_resolved")
            self.assertEqual(by_family["fam-rdasb-publication"]["locator_type"], "year")
            self.assertEqual(by_family["fam-plate-references"]["resolution_level"], "internal_reference")
            self.assertEqual(by_family["fam-plate-references"]["locator_type"], "plate")
            self.assertEqual(by_family["fam-plate-references"]["source_work_key"], "lucePeMaungTinInscriptionsOfBurma")
            self.assertEqual(by_family["fam-list-catalogue"]["bibtex_key"], "duroiselle1921list")
            self.assertEqual(by_family["fam-list-catalogue"]["locator_type"], "catalogue_number")
            self.assertEqual(by_family["fam-list-catalogue"]["source_work_key"], "duroiselle1921list")
            self.assertEqual(by_family["fam-raw-ub"]["source_family_id"], "sf-ub")
            self.assertEqual(by_family["fam-raw-ub"]["locator_type"], "volume_page")
            self.assertEqual(by_family["fam-raw-mp"]["source_family_id"], "sf-mp")
            self.assertEqual(by_family["fam-raw-mp"]["source_work_key"], "mandalayPalaceStoneCollection")
            self.assertEqual(by_family["fam-ppa-catalogue"]["source_family_id"], "sf-ppa")
            self.assertEqual(by_family["fam-ppa-catalogue"]["locator_type"], "page")
            self.assertEqual(by_family["fam-raw-a"]["source_family_id"], "sf-a")
            self.assertEqual(by_family["fam-raw-b"]["source_family_id"], "sf-b")
            self.assertIn("sf-rdasb", source_family_rows)
            self.assertIn("sf-ub", source_family_rows)
            self.assertEqual(acronym_rows["Pl."]["resolution_status"], "internal_locator")
            self.assertEqual(acronym_rows["IOB"]["resolution_status"], "internal_locator")
            self.assertEqual(acronym_rows["PPA"]["resolution_status"], "confirmed_expansion")
            self.assertEqual(acronym_rows["SIP"]["resolution_status"], "confirmed_expansion")
            self.assertEqual(
                acronym_rows["SIP"]["current_expansion"],
                "Pe Maung Tin and G. H. Luce, Selections from the Inscriptions of Pagan",
            )
            self.assertEqual(acronym_rows["UB"]["resolution_status"], "confirmed_expansion")
            self.assertEqual(acronym_rows["RDASB"]["resolution_status"], "probable_expansion")
            self.assertEqual(acronym_rows["MP"]["resolution_status"], "probable_locator_system")
            self.assertEqual(acronym_rows["OR"]["resolution_status"], "probable_locator_system")
            self.assertEqual(acronym_rows["Luce D"]["resolution_status"], "probable_private_luce_locator_system")
            self.assertEqual(acronym_rows["Luce J"]["resolution_status"], "probable_private_luce_locator_system")
            self.assertEqual(acronym_rows["IPPA"]["resolution_status"], "alias_or_variant_of_PPA")
            self.assertTrue(set(PRIORITY_ACRONYMS).issubset(acronym_rows))
            self.assertEqual(source_family_rows["sf-ppa"]["expanded_label"], "Inscriptions of Pagan, Pinya and Ava")
            self.assertEqual(source_family_rows["sf-mp"]["expanded_label"], "Mandalay Palace stone collection locator system")
            self.assertEqual(source_family_rows["sf-pl"]["expanded_label"], "Plate reference into Inscriptions of Burma")
            self.assertEqual(source_family_rows["sf-pl"]["locator_type"], "plate")
            self.assertEqual(source_family_rows["sf-pl"]["source_work_key"], "lucePeMaungTinInscriptionsOfBurma")
            self.assertEqual(len(remaining_worklist_rows), len(REMAINING_ACRONYMS))
            self.assertTrue(any(row["source_work_key"] == "lucePeMaungTinInscriptionsOfBurma" for row in locator_rows))
            self.assertTrue(any(row["source_work_key"] == "mandalayPalaceStoneCollection" for row in locator_rows))
            self.assertEqual(len(final_sprint_rows), 6)
            self.assertEqual(final_sprint_rows["RDASB"]["recommended_status"], "probable_expansion")
            self.assertEqual(final_sprint_rows["IPPA"]["recommended_status"], "alias_or_variant_of_PPA")
            self.assertEqual(unresolved_dossier_rows, [])
            self.assertTrue((output_dir / "ippa_occurrence_contexts.tsv").exists())
            self.assertTrue((output_dir / "ippa_ppa_comparison.tsv").exists())
            self.assertTrue((output_dir / "ippa_local_context_search.tsv").exists())
            self.assertTrue((output_dir / "ippa_frasch_abbrev_neighbourhood.tsv").exists())
            self.assertTrue((output_dir / "ippa_record_review.tsv").exists())
            self.assertTrue((output_dir / "ippa_targeted_ocr_notes.tsv").exists())
            self.assertTrue((output_dir / "ippa_resolution_decision.tsv").exists())
            self.assertNotIn("workUnresolved", candidate_bib)
            self.assertNotIn("RDASB", [row["family_label"] for row in report["top_unresolved_families"]])

    def test_parse_locator_handles_plate_iob_list_and_ppa_patterns(self) -> None:
        locator, locator_type = parse_locator("Pl. II 198", "fam-plate-references", "Pl.")
        self.assertEqual(locator, "II 198")
        self.assertEqual(locator_type, "plate")

        locator, locator_type = parse_locator("IOB--278", "fam-iob-catalogue", "IOB")
        self.assertEqual(locator, "278")
        self.assertEqual(locator_type, "catalogue_number")

        locator, locator_type = parse_locator("List 90", "fam-list-catalogue", "List")
        self.assertEqual(locator, "90")
        self.assertEqual(locator_type, "catalogue_number")

    def test_build_ippa_review_artifacts_extracts_all_occurrences_from_synthetic_corpus(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            corpus_path = temp_path / "inscriptions.jsonl"
            corpus_path.write_text(
                "".join(
                    json.dumps(row, ensure_ascii=False) + "\n"
                    for row in [
                        {
                            "record_id": "r1",
                            "source_deposit": "zenodo_4321314",
                            "source_layer": "structured_corpus",
                            "source_volume": "1",
                            "source_inscription_number": "4",
                            "source_page": "11",
                            "face": "ob",
                            "title_original": "Test one",
                            "references_original": "IPPA-159; PPA, p. 159; OBI 1, p. 11",
                        },
                        {
                            "record_id": "r2",
                            "source_deposit": "zenodo_4321314",
                            "source_layer": "structured_corpus",
                            "source_volume": "1",
                            "source_inscription_number": "5",
                            "source_page": "12",
                            "face": "ob",
                            "title_original": "Test two",
                            "references_original": "IPPA 137-138; List 90",
                        },
                    ]
                ),
                encoding="utf-8",
            )
            reference_occurrences_path = temp_path / "reference_occurrences.tsv"
            write_tsv(
                reference_occurrences_path,
                [
                    {"record_id": "r1", "raw_reference_string": "IPPA-159"},
                    {"record_id": "r2", "raw_reference_string": "IPPA 137-138"},
                ],
                ["record_id", "raw_reference_string"],
            )
            frasch_path = temp_path / "frasch_extracted_text.txt"
            frasch_path.write_text("PPA Arch. Survey of Burma (ed.), Inscriptions of Pagan, Pinya and Ava\n", encoding="utf-8")

            review = build_ippa_review_artifacts(
                corpus_inscriptions_path=corpus_path,
                reference_occurrences_path=reference_occurrences_path,
                frasch_extracted_text_path=frasch_path,
            )

            self.assertEqual(len(review["occurrence_rows"]), 2)
            self.assertEqual(review["occurrence_rows"][0]["record_id"], "r1")
            self.assertEqual(review["comparison_rows"][0]["looks_like_alias"], "true")
            self.assertEqual(review["decision_row"]["decision"], "alias_or_variant_of_PPA")
            self.assertEqual(review["decision_row"]["occurrences_reviewed"], "2")

    def test_build_authority_preserves_raw_ippa_strings_while_mapping_to_ppa(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            families_path, members_path, candidates_path, seeds_path, external_entries_path = self.write_fixture_tables(temp_path)
            acronym_candidates_path, acronym_report_path = self.write_fixture_acronym_files(temp_path)
            output_dir = temp_path / "authority"

            family_rows = read_tsv(families_path)
            family_rows.extend(
                [
                    {
                        "family_id": "fam-ippa-catalogue",
                        "family_label": "IPPA references",
                        "family_type": "source_catalogue",
                        "member_count": "1",
                        "occurrence_count": "1",
                        "sample_raw_references": "IPPA-159",
                        "likely_contains_translation": "no",
                        "review_status": "needs_human_review",
                        "notes": "",
                    },
                    {
                        "family_id": "fam-ppa-catalogue",
                        "family_label": "PPA references",
                        "family_type": "source_catalogue",
                        "member_count": "1",
                        "occurrence_count": "1",
                        "sample_raw_references": "PPA, p. 159",
                        "likely_contains_translation": "no",
                        "review_status": "needs_human_review",
                        "notes": "",
                    },
                ]
            )
            write_tsv(families_path, family_rows, FAMILY_FIELDS)

            member_rows = read_tsv(members_path)
            member_rows.extend(
                [
                    {
                        "family_id": "fam-ippa-catalogue",
                        "raw_reference_string": "IPPA-159",
                        "occurrence_count": "1",
                        "example_record_ids": "r1",
                        "notes": "",
                    },
                    {
                        "family_id": "fam-ppa-catalogue",
                        "raw_reference_string": "PPA, p. 159",
                        "occurrence_count": "1",
                        "example_record_ids": "r1",
                        "notes": "",
                    },
                ]
            )
            write_tsv(members_path, member_rows, MEMBER_FIELDS)

            build_authority(
                reference_families_path=families_path,
                reference_members_path=members_path,
                work_candidates_path=candidates_path,
                seed_path=seeds_path,
                external_entries_path=external_entries_path,
                output_dir=output_dir,
                frasch_references_path=temp_path / "missing_frasch.tsv",
                local_candidates_path=temp_path / "missing_local_candidates.tsv",
                local_manifest_path=temp_path / "missing_local_manifest.tsv",
                acronym_candidates_path=acronym_candidates_path,
                acronym_report_path=acronym_report_path,
            )

            crosswalk_rows = read_tsv(output_dir / "raw_reference_to_bibtex.tsv")
            ippa_rows = [row for row in crosswalk_rows if row["family_id"] == "fam-ippa-catalogue"]
            source_family_rows = {row["source_family_id"]: row for row in read_tsv(output_dir / "source_family_authority.tsv")}

            self.assertEqual(len(ippa_rows), 1)
            self.assertEqual(ippa_rows[0]["raw_reference_string"], "IPPA-159")
            self.assertEqual(ippa_rows[0]["source_family_id"], "sf-ippa")
            self.assertEqual(ippa_rows[0]["source_work_key"], "ppaCatalogue")
            self.assertEqual(ippa_rows[0]["bibtex_key"], "ppaCatalogue")
            self.assertEqual(source_family_rows["sf-ippa"]["alias_of_source_family_id"], "sf-ppa")

    def test_validate_bibtex_authority_requires_ippa_decision_artifact(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            families_path, members_path, candidates_path, seeds_path, external_entries_path = self.write_fixture_tables(temp_path)
            acronym_candidates_path, acronym_report_path = self.write_fixture_acronym_files(temp_path)
            output_dir = temp_path / "authority"

            build_authority(
                reference_families_path=families_path,
                reference_members_path=members_path,
                work_candidates_path=candidates_path,
                seed_path=seeds_path,
                external_entries_path=external_entries_path,
                output_dir=output_dir,
                frasch_references_path=temp_path / "missing_frasch.tsv",
                local_candidates_path=temp_path / "missing_local_candidates.tsv",
                local_manifest_path=temp_path / "missing_local_manifest.tsv",
                acronym_candidates_path=acronym_candidates_path,
                acronym_report_path=acronym_report_path,
            )

            (output_dir / "ippa_resolution_decision.tsv").unlink()

            result = validate_bibtex_authority(
                authority_bib_path=output_dir / "bibliography_authority.bib",
                candidates_bib_path=output_dir / "bibliography_candidates.bib",
                authority_tsv_path=output_dir / "bibtex_authority.tsv",
                crosswalk_path=output_dir / "raw_reference_to_bibtex.tsv",
                families_path=families_path,
                external_entries_path=external_entries_path,
                seed_path=seeds_path,
                high_frequency_path=output_dir / "high_frequency_unresolved.tsv",
                evidence_path=output_dir / "bibtex_authority_evidence.tsv",
                resolution_plan_path=output_dir / "high_frequency_resolution_plan.tsv",
                source_family_path=output_dir / "source_family_authority.tsv",
                report_path=output_dir / "bibtex_authority_report.json",
                frasch_references_path=temp_path / "missing_frasch.tsv",
                local_manifest_path=temp_path / "missing_local_manifest.tsv",
                acronym_status_path=output_dir / "acronym_resolution_status.tsv",
                acronym_candidates_path=acronym_candidates_path,
                acronym_report_path=acronym_report_path,
                manual_review_packet_path=output_dir / "acronym_manual_review_packet.tsv",
                remaining_acronym_worklist_path=output_dir / "remaining_acronym_worklist.tsv",
                remaining_acronym_evidence_path=output_dir / "remaining_acronym_evidence.tsv",
                source_work_locator_systems_path=output_dir / "source_work_locator_systems.tsv",
                source_work_authority_path=output_dir / "source_work_authority.tsv",
                raw_reference_crosswalk_audit_path=output_dir / "raw_reference_crosswalk_audit.tsv",
                candidate_stub_review_path=output_dir / "candidate_stub_review.tsv",
                final_acronym_resolution_sprint_path=output_dir / "final_acronym_resolution_sprint.tsv",
                final_acronym_local_file_hits_path=output_dir / "final_acronym_local_file_hits.tsv",
                final_acronym_web_searches_path=output_dir / "final_acronym_web_searches.tsv",
                frasch_abbreviation_list_review_path=output_dir / "frasch_abbreviation_list_review.tsv",
                unresolved_acronym_dossier_path=output_dir / "unresolved_acronym_dossier.tsv",
                ippa_occurrence_contexts_path=output_dir / "ippa_occurrence_contexts.tsv",
                ippa_ppa_comparison_path=output_dir / "ippa_ppa_comparison.tsv",
                ippa_local_context_search_path=output_dir / "ippa_local_context_search.tsv",
                ippa_frasch_abbrev_neighbourhood_path=output_dir / "ippa_frasch_abbrev_neighbourhood.tsv",
                ippa_record_review_path=output_dir / "ippa_record_review.tsv",
                ippa_targeted_ocr_notes_path=output_dir / "ippa_targeted_ocr_notes.tsv",
                ippa_resolution_decision_path=output_dir / "ippa_resolution_decision.tsv",
                reference_occurrences_path=ROOT / "data/working/bibliography/reference_occurrences.tsv",
            )

            self.assertFalse(result["ok"])
            self.assertTrue(any("ippa_resolution_decision.tsv is missing" in error for error in result["errors"]))

        locator, locator_type = parse_locator("PPA, p. 55", "fam-ppa-catalogue", "PPA")
        self.assertEqual(locator, "p. 55")
        self.assertEqual(locator_type, "page")

        locator, locator_type = parse_locator("OR 3434, fol. gha verso", "fam-raw-or", "OR")
        self.assertEqual(locator, "3434, fol. gha verso")
        self.assertEqual(locator_type, "folio")

        locator, locator_type = parse_locator("Luce D 825", "fam-raw-luce-d", "Luce D")
        self.assertEqual(locator, "825")
        self.assertEqual(locator_type, "number")

    def test_build_authority_applies_manual_acronym_seeds(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            families_path, members_path, candidates_path, seeds_path, external_entries_path = self.write_fixture_tables(temp_path)
            acronym_candidates_path, acronym_report_path = self.write_fixture_acronym_files(temp_path)
            manual_seed_path = self.write_manual_seed_file(temp_path)
            output_dir = temp_path / "authority"

            report = build_authority(
                reference_families_path=families_path,
                reference_members_path=members_path,
                work_candidates_path=candidates_path,
                seed_path=seeds_path,
                external_entries_path=external_entries_path,
                output_dir=output_dir,
                frasch_references_path=temp_path / "missing_frasch.tsv",
                local_candidates_path=temp_path / "missing_local_candidates.tsv",
                local_manifest_path=temp_path / "missing_local_manifest.tsv",
                acronym_candidates_path=acronym_candidates_path,
                acronym_report_path=acronym_report_path,
                manual_acronym_seeds_path=manual_seed_path,
            )

            acronym_rows = {row["acronym"]: row for row in read_tsv(output_dir / "acronym_resolution_status.tsv")}
            self.assertEqual(acronym_rows["EB"]["current_expansion"], "Epigraphia Birmanica")
            self.assertEqual(acronym_rows["JBRS"]["current_expansion"], "Journal of the Burma Research Society")
            self.assertEqual(acronym_rows["JRAS"]["current_expansion"], "Journal of the Royal Asiatic Society")
            self.assertEqual(acronym_rows["OBI"]["current_expansion"], "Old Burmese Inscriptions")
            self.assertEqual(acronym_rows["JBRS"]["definition_quality"], "manual_seed")
            self.assertEqual(acronym_rows["OBI"]["confidence"], "high")
            self.assertIn("Nathan", acronym_rows["OBI"]["notes"])
            self.assertEqual(report["manual_acronym_seed_count"], 4)
            self.assertEqual(report["manual_review_packet_rows"], len(PRIORITY_ACRONYMS))

    def test_validate_bibtex_authority_accepts_manual_seed_outputs(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            families_path, members_path, candidates_path, seeds_path, external_entries_path = self.write_fixture_tables(temp_path)
            acronym_candidates_path, acronym_report_path = self.write_fixture_acronym_files(temp_path)
            manual_seed_path = self.write_manual_seed_file(temp_path)
            output_dir = temp_path / "authority"

            build_authority(
                reference_families_path=families_path,
                reference_members_path=members_path,
                work_candidates_path=candidates_path,
                seed_path=seeds_path,
                external_entries_path=external_entries_path,
                output_dir=output_dir,
                frasch_references_path=temp_path / "missing_frasch.tsv",
                local_candidates_path=temp_path / "missing_local_candidates.tsv",
                local_manifest_path=temp_path / "missing_local_manifest.tsv",
                acronym_candidates_path=acronym_candidates_path,
                acronym_report_path=acronym_report_path,
                manual_acronym_seeds_path=manual_seed_path,
            )

            result = validate_bibtex_authority(
                authority_bib_path=output_dir / "bibliography_authority.bib",
                candidates_bib_path=output_dir / "bibliography_candidates.bib",
                authority_tsv_path=output_dir / "bibtex_authority.tsv",
                crosswalk_path=output_dir / "raw_reference_to_bibtex.tsv",
                families_path=families_path,
                external_entries_path=external_entries_path,
                seed_path=seeds_path,
                high_frequency_path=output_dir / "high_frequency_unresolved.tsv",
                evidence_path=output_dir / "bibtex_authority_evidence.tsv",
                resolution_plan_path=output_dir / "high_frequency_resolution_plan.tsv",
                source_family_path=output_dir / "source_family_authority.tsv",
                report_path=output_dir / "bibtex_authority_report.json",
                frasch_references_path=temp_path / "missing_frasch.tsv",
                local_manifest_path=temp_path / "missing_local_manifest.tsv",
                acronym_status_path=output_dir / "acronym_resolution_status.tsv",
                acronym_candidates_path=acronym_candidates_path,
                acronym_report_path=acronym_report_path,
                manual_acronym_seeds_path=manual_seed_path,
                manual_review_packet_path=output_dir / "acronym_manual_review_packet.tsv",
                remaining_acronym_worklist_path=output_dir / "remaining_acronym_worklist.tsv",
                remaining_acronym_evidence_path=output_dir / "remaining_acronym_evidence.tsv",
                source_work_locator_systems_path=output_dir / "source_work_locator_systems.tsv",
                source_work_authority_path=output_dir / "source_work_authority.tsv",
                raw_reference_crosswalk_audit_path=output_dir / "raw_reference_crosswalk_audit.tsv",
                candidate_stub_review_path=output_dir / "candidate_stub_review.tsv",
                final_acronym_resolution_sprint_path=output_dir / "final_acronym_resolution_sprint.tsv",
                final_acronym_local_file_hits_path=output_dir / "final_acronym_local_file_hits.tsv",
                final_acronym_web_searches_path=output_dir / "final_acronym_web_searches.tsv",
                frasch_abbreviation_list_review_path=output_dir / "frasch_abbreviation_list_review.tsv",
                unresolved_acronym_dossier_path=output_dir / "unresolved_acronym_dossier.tsv",
            )
            self.assertTrue(result["ok"], result["errors"])

    def test_validate_bibtex_authority_rejects_missing_final_web_search_record(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            families_path, members_path, candidates_path, seeds_path, external_entries_path = self.write_fixture_tables(temp_path)
            acronym_candidates_path, acronym_report_path = self.write_fixture_acronym_files(temp_path)
            output_dir = temp_path / "authority"

            build_authority(
                reference_families_path=families_path,
                reference_members_path=members_path,
                work_candidates_path=candidates_path,
                seed_path=seeds_path,
                external_entries_path=external_entries_path,
                output_dir=output_dir,
                frasch_references_path=temp_path / "missing_frasch.tsv",
                local_candidates_path=temp_path / "missing_local_candidates.tsv",
                local_manifest_path=temp_path / "missing_local_manifest.tsv",
                acronym_candidates_path=acronym_candidates_path,
                acronym_report_path=acronym_report_path,
            )

            web_search_path = output_dir / "final_acronym_web_searches.tsv"
            web_rows = [row for row in read_tsv(web_search_path) if row["acronym"] != "IPPA"]
            write_tsv(web_search_path, web_rows, [
                "acronym",
                "query",
                "result_title",
                "result_url_or_domain",
                "short_result_summary",
                "supports_candidate_expansion",
                "confidence",
                "notes",
            ])

            result = validate_bibtex_authority(
                authority_bib_path=output_dir / "bibliography_authority.bib",
                candidates_bib_path=output_dir / "bibliography_candidates.bib",
                authority_tsv_path=output_dir / "bibtex_authority.tsv",
                crosswalk_path=output_dir / "raw_reference_to_bibtex.tsv",
                families_path=families_path,
                external_entries_path=external_entries_path,
                seed_path=seeds_path,
                high_frequency_path=output_dir / "high_frequency_unresolved.tsv",
                evidence_path=output_dir / "bibtex_authority_evidence.tsv",
                resolution_plan_path=output_dir / "high_frequency_resolution_plan.tsv",
                source_family_path=output_dir / "source_family_authority.tsv",
                report_path=output_dir / "bibtex_authority_report.json",
                frasch_references_path=temp_path / "missing_frasch.tsv",
                local_manifest_path=temp_path / "missing_local_manifest.tsv",
                acronym_status_path=output_dir / "acronym_resolution_status.tsv",
                acronym_candidates_path=acronym_candidates_path,
                acronym_report_path=acronym_report_path,
                manual_review_packet_path=output_dir / "acronym_manual_review_packet.tsv",
                remaining_acronym_worklist_path=output_dir / "remaining_acronym_worklist.tsv",
                remaining_acronym_evidence_path=output_dir / "remaining_acronym_evidence.tsv",
                source_work_locator_systems_path=output_dir / "source_work_locator_systems.tsv",
                source_work_authority_path=output_dir / "source_work_authority.tsv",
                raw_reference_crosswalk_audit_path=output_dir / "raw_reference_crosswalk_audit.tsv",
                candidate_stub_review_path=output_dir / "candidate_stub_review.tsv",
                final_acronym_resolution_sprint_path=output_dir / "final_acronym_resolution_sprint.tsv",
                final_acronym_local_file_hits_path=output_dir / "final_acronym_local_file_hits.tsv",
                final_acronym_web_searches_path=web_search_path,
                frasch_abbreviation_list_review_path=output_dir / "frasch_abbreviation_list_review.tsv",
                unresolved_acronym_dossier_path=output_dir / "unresolved_acronym_dossier.tsv",
            )
            self.assertFalse(result["ok"])
            self.assertTrue(any("final_acronym_web_searches.tsv is missing web-search rows for IPPA" in error for error in result["errors"]))

    def test_build_authority_marks_final_sprint_acronyms_with_precise_statuses(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            families_path, members_path, candidates_path, seeds_path, external_entries_path = self.write_fixture_tables(temp_path)
            acronym_candidates_path, acronym_report_path = self.write_fixture_acronym_files(temp_path)
            output_dir = temp_path / "authority"

            build_authority(
                reference_families_path=families_path,
                reference_members_path=members_path,
                work_candidates_path=candidates_path,
                seed_path=seeds_path,
                external_entries_path=external_entries_path,
                output_dir=output_dir,
                frasch_references_path=temp_path / "missing_frasch.tsv",
                local_candidates_path=temp_path / "missing_local_candidates.tsv",
                local_manifest_path=temp_path / "missing_local_manifest.tsv",
                acronym_candidates_path=acronym_candidates_path,
                acronym_report_path=acronym_report_path,
            )

            acronym_rows = {row["acronym"]: row for row in read_tsv(output_dir / "acronym_resolution_status.tsv")}
            self.assertEqual(acronym_rows["IPPA"]["resolution_status"], "alias_or_variant_of_PPA")
            self.assertEqual(acronym_rows["RDASB"]["resolution_status"], "probable_expansion")
            self.assertEqual(acronym_rows["MP"]["resolution_status"], "probable_locator_system")
            self.assertEqual(acronym_rows["OR"]["resolution_status"], "probable_locator_system")
            self.assertEqual(acronym_rows["Luce D"]["resolution_status"], "probable_private_luce_locator_system")
            self.assertEqual(acronym_rows["Luce J"]["resolution_status"], "probable_private_luce_locator_system")
            self.assertEqual(acronym_rows["TN"]["resolution_status"], "confirmed_expansion")
            self.assertEqual(acronym_rows["UEM"]["resolution_status"], "confirmed_expansion")

    def test_extract_documentation_sections_rejects_generic_and_irrelevant_bibliography(self) -> None:
        generic_rows = extract_documentation_sections(
            "Bibliography\nA. Author, Some Book.\nB. Author, Another Book.\n",
            source_file_id="doc-1",
            source_file_label="generic bibliography.txt",
        )
        tibetan_rows = extract_documentation_sections(
            "Bibliography\nRichardson, Tibetan Inscriptions. OBI appears only as a parenthetical note.\n",
            source_file_id="doc-2",
            source_file_label="Richardson Tibetan corpus bibliography.pdf",
        )
        self.assertEqual(generic_rows, [])
        self.assertEqual(tibetan_rows, [])

    def test_ocr_priority_sources_writes_manifest_and_snippet_index_for_mock_text(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            source_root = temp_path / "sources"
            source_dir = source_root / "mock-source-123"
            source_dir.mkdir(parents=True)
            source_path = source_dir / "mock-source.txt"
            full_text = (
                "Introductory material about Burmese epigraphy and corpus notes.\n"
                "Abbreviations\n"
                "SIP = Pe Maung Tin and G. H. Luce, Selections from the Inscriptions of Pagan\n"
                "JBRS = Journal of the Burma Research Society\n"
                "More explanatory notes.\n"
            ) * 6
            source_path.write_text(full_text, encoding="utf-8")

            queue_path = temp_path / "ocr_priority_queue.tsv"
            write_tsv(
                queue_path,
                [
                    {
                        "source_file_id": "mock-source-123",
                        "source_file_label": "mock-source.txt",
                        "reason_ocr_needed": "mock test",
                        "priority": "high",
                        "priority_reason": "test",
                        "target_acronyms": "SIP, JBRS",
                        "expected_value": "abbreviation definitions",
                        "notes": "",
                    }
                ],
                ["source_file_id", "source_file_label", "reason_ocr_needed", "priority", "priority_reason", "target_acronyms", "expected_value", "notes"],
            )
            output_dir = temp_path / "ocr_outputs"
            local_text_root = temp_path / "local_ocr_text"

            subprocess.run(
                [
                    sys.executable,
                    str(SCRIPTS_DIR / "ocr_priority_sources.py"),
                    "--queue",
                    str(queue_path),
                    "--source-root",
                    str(source_root),
                    "--manifest",
                    str(temp_path / "missing_manifest.tsv"),
                    "--output-dir",
                    str(output_dir),
                    "--local-text-root",
                    str(local_text_root),
                ],
                check=True,
            )

            manifest_rows = read_tsv(output_dir / "ocr_manifest.tsv")
            index_rows = read_tsv(output_dir / "ocr_text_index.tsv")
            self.assertEqual(manifest_rows[0]["extraction_status"], "success")
            self.assertTrue((local_text_root / "mock-source-123.txt").exists())
            self.assertTrue(any("SIP" in row["snippet_text"] for row in index_rows))
            self.assertFalse(any(path.suffix == ".txt" for path in output_dir.iterdir()))

    def test_validate_bibtex_authority_detects_duplicate_key(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            families_path = temp_path / "reference_families.tsv"
            crosswalk_path = temp_path / "raw_reference_to_bibtex.tsv"
            authority_tsv_path = temp_path / "bibtex_authority.tsv"
            authority_bib_path = temp_path / "bibliography_authority.bib"
            candidate_bib_path = temp_path / "bibliography_candidates.bib"
            external_entries_path = temp_path / "external.tsv"

            write_tsv(families_path, [{"family_id": "fam-1", "family_label": "OBI", "family_type": "source_catalogue", "member_count": "1", "occurrence_count": "1", "sample_raw_references": "OBI 1", "likely_contains_translation": "no", "review_status": "unreviewed", "notes": ""}], FAMILY_FIELDS)
            write_tsv(
                authority_tsv_path,
                [
                    {
                        "bibtex_key": "dupKey",
                        "entry_type": "misc",
                        "authority_status": "machine_stub",
                        "source_of_authority": "corpus_reference",
                        "matched_external_key": "",
                        "resolution_status": "needs_human_review",
                        "resolution_level": "unknown",
                        "source_family_id": "",
                        "family_id": "fam-1",
                        "family_label": "OBI",
                        "family_type": "source_catalogue",
                        "author": "",
                        "editor": "",
                        "year": "",
                        "title": "OBI",
                        "shorttitle": "OBI",
                        "journal": "",
                        "booktitle": "",
                        "publisher": "",
                        "address": "",
                        "volume": "",
                        "number": "",
                        "pages": "",
                        "doi": "",
                        "url": "",
                        "isbn": "",
                        "language": "",
                        "script": "",
                        "translation_relevance": "unknown",
                        "review_status": "unreviewed",
                        "evidence": "test",
                        "notes": "test",
                    }
                ],
                AUTHORITY_FIELDS,
            )
            write_tsv(
                crosswalk_path,
                [
                    {
                        "raw_reference_string": "OBI 1",
                        "family_id": "fam-1",
                        "source_family_id": "",
                        "work_candidate_id": "wc-1",
                        "bibtex_key": "dupKey",
                        "locator": "1",
                        "locator_type": "number",
                        "resolution_status": "needs_human_review",
                        "resolution_level": "unknown",
                        "match_type": "machine_stub_match",
                        "match_confidence": "low",
                        "evidence": "test",
                        "needs_human_review": "true",
                        "notes": "",
                    }
                ],
                CROSSWALK_FIELDS,
            )
            write_tsv(external_entries_path, [], ["bibtex_key"])
            authority_bib_path.write_text("@misc{dupKey,\n  title = {One},\n}\n", encoding="utf-8")
            candidate_bib_path.write_text("@misc{dupKey,\n  title = {Two},\n}\n", encoding="utf-8")

            result = validate_bibtex_authority(
                authority_bib_path=authority_bib_path,
                candidates_bib_path=candidate_bib_path,
                authority_tsv_path=authority_tsv_path,
                crosswalk_path=crosswalk_path,
                families_path=families_path,
                external_entries_path=external_entries_path,
            )
            self.assertFalse(result["ok"])
            self.assertTrue(any("duplicate BibTeX keys" in error for error in result["errors"]))


if __name__ == "__main__":
    unittest.main()
